"""文档导出 - 生成带颜色标注的 .docx"""
import io
import difflib
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX


def generate_diff_docx(titles, text_before, text_after, tags, title_label="审稿对比"):
    """生成带红绿黄标注的对比 .docx 文件

    标注规则：
    - 红色划线：删除的内容
    - 黄底划线：被替换的原文
    - 绿色底色：新增/替换后的内容

    Returns: BytesIO
    """
    doc = Document()
    doc.add_heading(title_label, level=1)

    # ── 标题 ──
    for i, t in enumerate(titles):
        p = doc.add_paragraph()
        run = p.add_run(f"标题{i+1}：{t}")
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph()

    # ── 正文（diff 标注）──
    doc.add_heading("正文（标注版）", level=2)
    sm = difflib.SequenceMatcher(None, text_before, text_after, autojunk=False)

    segments = []
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            segments.append(("normal", text_after[j1:j2]))
        elif op == "delete":
            segments.append(("deleted", text_before[i1:i2]))
        elif op == "insert":
            segments.append(("added", text_after[j1:j2]))
        elif op == "replace":
            segments.append(("replaced_old", text_before[i1:i2]))
            segments.append(("replaced_new", text_after[j1:j2]))

    p = doc.add_paragraph()
    for seg_type, text in segments:
        lines = text.split("\n")
        for li, line in enumerate(lines):
            if line:
                run = p.add_run(line)
                run.font.size = Pt(11)
                if seg_type == "deleted":
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                    run.font.strike = True
                elif seg_type == "added":
                    run.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                elif seg_type == "replaced_old":
                    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
                    run.font.strike = True
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif seg_type == "replaced_new":
                    run.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            if li < len(lines) - 1:
                p = doc.add_paragraph()

    # ── 标签 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("话题标签：")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(tags)
    run.font.size = Pt(11)

    # ── 图例 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("【标注说明】")
    run.bold = True
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run("红色划线")
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    run.font.strike = True
    run.font.size = Pt(9)
    run = p.add_run(" = 删除    ")
    run.font.size = Pt(9)

    run = p.add_run("黄底划线")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.strike = True
    run.font.size = Pt(9)
    run = p.add_run(" = 被替换原文    ")
    run.font.size = Pt(9)

    run = p.add_run("绿色底色")
    run.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    run.font.size = Pt(9)
    run = p.add_run(" = 新增/替换后")
    run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_clean_docx(titles, body, tags):
    """生成干净的终稿 .docx（无标注）

    Returns: BytesIO
    """
    doc = Document()
    doc.add_heading("终稿", level=1)

    for i, t in enumerate(titles):
        p = doc.add_paragraph()
        run = p.add_run(f"标题{i+1}：{t}")
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph()

    for line in body.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(tags)
    run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
