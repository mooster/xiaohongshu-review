import streamlit as st
import re
import os
import json
from datetime import datetime
from docx import Document
import io
import urllib.request

TODAY = datetime.now().strftime("%Y%m%d")

# ========== å®¡æ ¸è§„åˆ™å¸¸é‡ ==========
REQUIRED_TAGS = [
    "#èƒ½æ©å…¨æŠ¤", "#èƒ½æ©å…¨æŠ¤æ°´å¥¶", "#é€‚åº¦æ°´è§£", "#é€‚åº¦æ°´è§£å¥¶ç²‰",
    "#é€‚åº¦æ°´è§£å¥¶ç²‰æ¨è", "#é˜²æ•å¥¶ç²‰", "#ç¬¬ä¸€å£å¥¶ç²‰", "#é›€å·¢é€‚åº¦æ°´è§£"
]

TITLE_KEYWORDS = ["é€‚åº¦æ°´è§£", "é˜²æ•", "ç§‘æ™®"]
BODY_KEYWORDS = ["é€‚åº¦æ°´è§£", "é˜²æ•", "èƒ½æ©å…¨æŠ¤"]
COVER_KEYWORDS = ["é€‚åº¦æ°´è§£", "é˜²æ•", "ç§‘æ™®"]

FORBIDDEN_WORDS = {
    "ç¦æ­¢è¯": ["æ•å®", "å¥¶ç“¶", "å¥¶å˜´", "æ–°ç”Ÿå„¿", "è¿‡æ•", "ç–¾ç—…"],
    "ç¦ç–—æ•ˆè¡¨è¿°": ["é¢„é˜²", "ç”Ÿé•¿", "å‘è‚²", "å…ç–«"],
    "ç¦ç»å¯¹åŒ–": ["æœ€", "ç¬¬ä¸€", "TOP1"],
}

FORBIDDEN_EXCEPTIONS = {
    "ç¬¬ä¸€": ["ç¬¬ä¸€å£å¥¶ç²‰", "ç¬¬ä¸€å£é…æ–¹ç²‰"],
    "æœ€": ["æœ€è¿‘", "æœ€å", "æœ€ç»ˆ", "æœ€åˆ", "æœ€å¤š"],
}

FORBIDDEN_REPLACEMENTS = {
    "è¿‡æ•": "æ•æ•", "æ•å®": "æ•æ„Ÿä½“è´¨å®å®",
    "æ–°ç”Ÿå„¿": "åˆç”Ÿå®å®", "é¢„é˜²": "é˜²æ•",
    "ç”Ÿé•¿": "æˆé•¿", "å‘è‚²": "å™Œå™Œé•¿",
    "å…ç–«": "ä¿æŠ¤åŠ›", "ç–¾ç—…": "ä¸é€‚",
}

# å¿…æéœ€æ¶¦è‰²å–ç‚¹ (4å¤§æ–¹å‘10å°æ–¹å‘)
PARAPHRASE_SELLING_POINTS = [
    {"category": "æ•æ•èƒŒæ™¯", "idx": 1,
     "text": "æˆ‘å›½åˆç”Ÿå®å®æ•æ•ç‡é«˜è¾¾40%ï¼Œè¦æ˜¯æœ‰çˆ¶æ¯æ•æ•å²ï¼Œå®å®æ•æ•çš„æ¦‚ç‡å°†é£™å‡åˆ°80%",
     "fragment": "æ•æ•ç‡é«˜è¾¾40%"},
    {"category": "é˜²æ•-æ°´è§£æŠ€æœ¯", "idx": 2,
     "text": "æ˜“æ•çš„å¤§åˆ†å­ç‰›å¥¶è›‹ç™½åˆ‡å‰²æˆæ¸©å’Œçš„é€‚åº¦æ°´è§£å°åˆ†å­ç‰›å¥¶è›‹ç™½ï¼Œç²¾å‡†å»æ‰è‡´æ•ç‰‡æ®µçš„åŒæ—¶ï¼Œåˆå®Œæ•´ä¿ç•™äº†è›‹ç™½æœ‰ç›Šè¥å…»",
     "fragment": "åˆ‡å‰²æˆæ¸©å’Œçš„é€‚åº¦æ°´è§£å°åˆ†å­"},
    {"category": "é˜²æ•-æ°´è§£æŠ€æœ¯", "idx": 3,
     "text": "å…¨çƒä¸“ä¸šäººå£«ä¼˜å…ˆæ¨èå‘¢",
     "fragment": "å…¨çƒä¸“ä¸šäººå£«ä¼˜å…ˆæ¨è"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 4,
     "text": "6ç§HMOåŠ ä¸Šæ˜æ˜ŸåŒèŒB.Infantis å’Œ Bb-12ï¼Œä¸¤è€…å¼ºå¼ºè”åˆï¼ŒååŒä½œç”¨é‡Šæ”¾é«˜å€çš„åŸç”Ÿä¿æŠ¤åŠ›",
     "fragment": "ä¸¤è€…å¼ºå¼ºè”åˆ"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 5,
     "text": "çŸ­çŸ­28å¤©å°±èƒ½è°ƒç†å¥½å¨ƒçš„è‚šè‚šèŒèŒç¯å¢ƒï¼Œä»è‚šè‚šåˆ°å…¨èº«éƒ½å»ºèµ·åšå›ºçš„é˜²æŠ¤å±éšœ",
     "fragment": "ä»è‚šè‚šåˆ°å…¨èº«"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 6,
     "text": "ä¿æŠ¤åŠ›èƒ½æŒç»­15ä¸ªæœˆï¼ŒåŠ©åŠ›å¨ƒæˆé•¿",
     "fragment": "åŠ©åŠ›å¨ƒæˆé•¿"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 7,
     "text": "å››ç»´æˆé•¿æ›²çº¿ç‰¹åˆ«å‡ºè‰²",
     "fragment": "å››ç»´æˆé•¿æ›²çº¿"},
    {"category": "åŸºç¡€è¥å…»", "idx": 8,
     "text": "åŸºç¡€è¥å…»ä¹Ÿå¾ˆæŠ—æ‰“",
     "fragment": "åŸºç¡€è¥å…»ä¹Ÿå¾ˆæŠ—æ‰“"},
    {"category": "åŸºç¡€è¥å…»", "idx": 9,
     "text": "25ç§ç»´ç”Ÿç´ å’ŒçŸ¿ç‰©è´¨æ‹‰æ»¡",
     "fragment": "ç»´ç”Ÿç´ å’ŒçŸ¿ç‰©è´¨æ‹‰æ»¡"},
    {"category": "åŸºç¡€è¥å…»", "idx": 10,
     "text": "å…¨ä¹³ç³–çš„é…æ–¹å£å‘³æ¸…æ·¡ï¼Œå®å®çˆ±å–",
     "fragment": "å…¨ä¹³ç³–çš„é…æ–¹å£å‘³æ¸…æ·¡ï¼Œå®å®çˆ±å–"},
]

# å¿…æä¸å¯ä¿®æ”¹å–ç‚¹ (3å¤§åˆ‡è§’10å°åˆ‡è§’)
FIXED_SELLING_POINTS = [
    {"category": "é˜²æ•-æ°´è§£æŠ€æœ¯", "idx": 1, "text": "å¤šé¡¹ç§‘å­¦å®è¯çš„é›€å·¢å°–å³°æ°´è§£æŠ€æœ¯"},
    {"category": "é˜²æ•-æ°´è§£æŠ€æœ¯", "idx": 2, "text": "æ¸©å’Œçš„é€‚åº¦æ°´è§£å°åˆ†å­ç‰›å¥¶è›‹ç™½"},
    {"category": "é˜²æ•-æ°´è§£æŠ€æœ¯", "idx": 3, "text": "é˜²æ•é¢†åŸŸæƒå¨å¾·å›½GINIç ”ç©¶è®¤è¯ï¼Œèƒ½é•¿æ•ˆé˜²æ•20å¹´ï¼Œç›¸æ¯”äºç‰›å¥¶è›‹ç™½è‡´æ•æ€§é™ä½1000å€"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 4, "text": "é‡‡ç”¨äº†å…¨çƒåˆ›æ–°çš„è¶…å€è‡ªæŠ¤ç§‘æŠ€"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 5, "text": "6ç§HMOåŠ ä¸Šæ˜æ˜ŸåŒèŒB.Infantis å’Œ Bb-12"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 6, "text": "ååŒä½œç”¨é‡Šæ”¾é«˜å€çš„åŸç”Ÿä¿æŠ¤åŠ›"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 7, "text": "çŸ­çŸ­28å¤©å°±èƒ½è°ƒç†å¥½å¨ƒçš„è‚šè‚šèŒèŒç¯å¢ƒ"},
    {"category": "è‡ªæŠ¤åŠ›", "idx": 8, "text": "ä¿æŠ¤åŠ›èƒ½æŒç»­15ä¸ªæœˆ"},
    {"category": "åŸºç¡€è¥å…»", "idx": 9, "text": "25ç§ç»´ç”Ÿç´ å’ŒçŸ¿ç‰©è´¨"},
    {"category": "åŸºç¡€è¥å…»", "idx": 10, "text": "å…¨ä¹³ç³–çš„é…æ–¹å£å‘³æ¸…æ·¡"},
]

# å…è®¸åˆ å‡çš„å–ç‚¹
OPTIONAL_SELLING_POINTS = [
    {"category": "é˜²æ•-æ°´è§£æŠ€æœ¯",
     "text": "æ¬§ç›Ÿè®¤å¯åŠå…¨çƒ30+ç§‘å­¦å®è¯èƒŒä¹¦ï¼Œç¡¬å®åŠ›çœŸæå®æ–™æ‘†å‡ºæ¥!",
     "fragment": "æ¬§ç›Ÿè®¤å¯"},
    {"category": "åŸºç¡€è¥å…»",
     "text": "æ—©æœŸé…æ–¹è¿˜æ­é…äº†ç‰›ç£ºé…¸ã€èƒ†ç¢±ã€æ ¸è‹·é…¸ç­‰å…³é”®è¥å…»ã€‚ä¸æ·»åŠ è”—ã€é¦™ç²¾è¿™äº›ä¸å‹å¥½æˆåˆ†ã€‚",
     "fragment": "ç‰›ç£ºé…¸ã€èƒ†ç¢±ã€æ ¸è‹·é…¸"},
]

# å–ç‚¹é¡ºåºé”šç‚¹
ORDER_ANCHORS = {
    "é˜²æ•-æ°´è§£æŠ€æœ¯": ["æ°´è§£æŠ€æœ¯", "å°–å³°æ°´è§£", "GINI", "è‡´æ•æ€§é™ä½", "é€‚åº¦æ°´è§£å°åˆ†å­"],
    "è‡ªæŠ¤åŠ›": ["è¶…å€è‡ªæŠ¤", "HMO", "åŒèŒ", "B.Infantis", "åŸç”Ÿä¿æŠ¤åŠ›"],
    "åŸºç¡€è¥å…»": ["ç»´ç”Ÿç´ å’ŒçŸ¿ç‰©è´¨", "å…¨ä¹³ç³–"],
}

# æ ‡å‡†å–ç‚¹ç¤ºä¾‹
SELLING_POINT_EXAMPLE = """æˆ‘å›½åˆç”Ÿå®å®æ•æ•ç‡é«˜è¾¾40%ï¼Œè¦æ˜¯æœ‰çˆ¶æ¯æ•æ•å²ï¼Œå®å®æ•æ•çš„æ¦‚ç‡å°†é£™å‡åˆ°80%ï¼›æ•æ•é«˜å‘çš„åŸå› ï¼ˆæœªç»äº§é“æŒ¤å‹ã€å…»å® ä¸“ä¸šäººå£«å»ºè®®ï¼šä¸å°‘ä¸“ä¸šäººå£«å»ºè®®ï¼Œå¯ä»¥ç»™å®å®é€‰æ‹©é€‚åº¦æ°´è§£é…æ–¹ç²‰ä½œä¸ºå®å®çš„ç¬¬ä¸€å£é…æ–¹ç²‰

æ‹¥æœ‰å¤šé¡¹ç§‘å­¦å®è¯çš„é›€å·¢å°–å³°æ°´è§£ç§‘æŠ€ï¼Œå°±åƒç»™è›‹ç™½è£…äº†ç²¾å‡†åˆ‡å‰²å™¨ï¼ŒæŠŠæ˜“æ•çš„å¤§åˆ†å­ç‰›å¥¶è›‹ç™½åˆ‡å‰²æˆæ¸©å’Œçš„100%é€‚åº¦æ°´è§£å°åˆ†å­ç‰›å¥¶è›‹ç™½ï¼Œç²¾å‡†å»æ‰è‡´æ•ç‰‡æ®µçš„åŒæ—¶ï¼Œåˆå®Œæ•´ä¿ç•™äº†è›‹ç™½æœ‰ç›Šè¥å…»ï¼Œæ›´äº²å’Œå®å®å¨‡è‚šè‚š!ä¸ä»…æœ‰é˜²æ•é¢†åŸŸæƒå¨å¾·å›½GINIç ”ç©¶è®¤è¯ï¼Œèƒ½é•¿æ•ˆé˜²æ•20å¹´ï¼Œè¿˜æœ‰æ¬§ç›Ÿè®¤å¯åŠå…¨çƒ30+ç§‘å­¦å®è¯èƒŒä¹¦ï¼Œç›¸æ¯”äºç‰›å¥¶è›‹ç™½è‡´æ•æ€§é™ä½1000å€ï¼Œç¡¬å®åŠ›çœŸæå®æ–™æ‘†å‡ºæ¥!æ€ªä¸å¾—å…¨çƒä¸“ä¸šäººå£«ä¼˜å…ˆæ¨èå‘¢ï¼/å…¨çƒä¸“ä¸šäººå£«ä¼˜å…ˆæ¨èæ˜¯æœ‰é“ç†çš„

å®ƒé‡‡ç”¨äº†å…¨çƒåˆ›æ–°çš„"è¶…å€è‡ªæŠ¤ç§‘æŠ€"ï¼Œå…¶ä¸­6ç§HMOåŠ ä¸Šæ˜æ˜ŸåŒèŒB.Infantis å’Œ Bb-12ï¼Œä¸¤è€…å¼ºå¼ºè”åˆï¼ŒååŒä½œç”¨é‡Šæ”¾é«˜å€çš„åŸç”Ÿä¿æŠ¤åŠ›ï¼çŸ­çŸ­28å¤©å°±èƒ½è°ƒç†å¥½å¨ƒçš„è‚šè‚šèŒèŒç¯å¢ƒï¼Œä»è‚šè‚šåˆ°å…¨èº«éƒ½å»ºèµ·åšå›ºçš„é˜²æŠ¤å±éšœã€‚æ›´å…³é”®çš„æ˜¯ï¼Œè¿™ä»½ä¿æŠ¤åŠ›èƒ½æŒç»­15ä¸ªæœˆï¼Œå®Œç¾è¦†ç›–å®å®çš„é»„é‡‘å‘è‚²æœŸï¼ŒåŠ©åŠ›å¨ƒå™Œé•¿ã€ç¨³ç¨³é•¿~ æœ‰å®ƒåŠ©åŠ›ï¼Œå¨ƒçš„å››ç»´æˆé•¿æ›²çº¿ç‰¹åˆ«å‡ºè‰²!

åŸºç¡€è¥å…»ä¹Ÿå¾ˆæŠ—æ‰“ï¼Œ25ç§ç»´ç”Ÿç´ å’ŒçŸ¿ç‰©è´¨æ‹‰æ»¡ï¼Œæ—©æœŸé…æ–¹è¿˜æ­é…äº†ç‰›ç£ºé…¸ã€èƒ†ç¢±ã€æ ¸è‹·é…¸ç­‰å…³é”®è¥å…»ã€‚å…¨ä¹³ç³–çš„é…æ–¹å£å‘³æ¸…æ·¡ï¼Œä¸æ·»åŠ è”—ã€é¦™ç²¾è¿™äº›ä¸å‹å¥½æˆåˆ†ï¼Œå®å®çˆ±å–ï¼Œå¦ˆå¦ˆæ”¾å¿ƒã€‚"""

# äººè¯ä¿®æ”¹ Prompt
RENHUA_PROMPT = """ä½ æ˜¯å°çº¢ä¹¦é¡¶çº§çˆ†æ–‡å†™æ‰‹ï¼Œæ“…é•¿æŠŠç¡¬å¹¿å†™æˆçœŸå®åˆ†äº«ã€‚ç°åœ¨å¸®æˆ‘æ”¹å†™ã€èƒ½æ©å…¨æŠ¤å¥¶ç²‰ã€‘çš„KOLç¨¿ä»¶ã€‚

âš ï¸ ã€æœ€é‡è¦çš„3ä¸ªç¡¬æ€§è¦æ±‚ - å¿…é¡»å…¨éƒ¨æ»¡è¶³ã€‘âš ï¸
1. æ­£æ–‡å­—æ•°å¿…é¡»åœ¨800-900å­—ä¹‹é—´ï¼ˆè¿™æ˜¯æœ€é‡è¦çš„ï¼å¤ªçŸ­æˆ–å¤ªé•¿éƒ½ä¸è¡Œï¼‰
2. å¿…é¡»æœ‰å¼ºçƒˆçš„å°çº¢ä¹¦æ´»äººæ„Ÿã€çˆ†æ–‡æ„Ÿã€çœŸå®åˆ†äº«æ„Ÿ
3. å¿…é¡»åŒ…å«ä¸‹é¢10å¥è¯æœ¯ï¼ˆå¯ä»¥è‡ªç„¶èå…¥ï¼Œä½†å­—å­—ä¸èƒ½æ”¹ï¼‰

ã€10å¥å¿…é¡»åŸå°ä¸åŠ¨å‡ºç°çš„è¯æœ¯ã€‘
â‘  å¤šé¡¹ç§‘å­¦å®è¯çš„é›€å·¢å°–å³°æ°´è§£æŠ€æœ¯
â‘¡ æ¸©å’Œçš„é€‚åº¦æ°´è§£å°åˆ†å­ç‰›å¥¶è›‹ç™½
â‘¢ é˜²æ•é¢†åŸŸæƒå¨å¾·å›½GINIç ”ç©¶è®¤è¯ï¼Œèƒ½é•¿æ•ˆé˜²æ•20å¹´ï¼Œç›¸æ¯”äºç‰›å¥¶è›‹ç™½è‡´æ•æ€§é™ä½1000å€
â‘£ é‡‡ç”¨äº†å…¨çƒåˆ›æ–°çš„è¶…å€è‡ªæŠ¤ç§‘æŠ€
â‘¤ 6ç§HMOåŠ ä¸Šæ˜æ˜ŸåŒèŒB.Infantis å’Œ Bb-12
â‘¥ ååŒä½œç”¨é‡Šæ”¾é«˜å€çš„åŸç”Ÿä¿æŠ¤åŠ›
â‘¦ çŸ­çŸ­28å¤©å°±èƒ½è°ƒç†å¥½å¨ƒçš„è‚šè‚šèŒèŒç¯å¢ƒ
â‘§ ä¿æŠ¤åŠ›èƒ½æŒç»­15ä¸ªæœˆ
â‘¨ 25ç§ç»´ç”Ÿç´ å’ŒçŸ¿ç‰©è´¨
â‘© å…¨ä¹³ç³–çš„é…æ–¹å£å‘³æ¸…æ·¡

ã€å°çº¢ä¹¦çˆ†æ–‡å†™æ³• - è¿™æ‰æ˜¯æ´»äººæ„Ÿï¼ã€‘
ğŸ”¥ å¼€å¤´è¦ç‚¸ï¼šç”¨"å§å¦¹ä»¬ï¼""æ•‘å‘½ï¼""åæ‚”æ²¡æ—©çŸ¥é“ï¼"ç­‰æƒ…ç»ªé’©å­å¼€åœº
ğŸ”¥ è¯´äººè¯ï¼šæŠŠ"å› æ­¤å»ºè®®"æ¢æˆ"æ‰€ä»¥æˆ‘çœŸå¿ƒæ¨è"ï¼ŒæŠŠ"å…·æœ‰"æ¢æˆ"æœ‰"
ğŸ”¥ åƒèŠå¤©ï¼šå¤šç”¨"æˆ‘""ä½ ""å’±å®¶å¨ƒ"ï¼Œå†™å¾—åƒåœ¨è·Ÿé—ºèœœåˆ†äº«ç»éªŒ
ğŸ”¥ æœ‰æƒ…ç»ªï¼šåŠ å…¥"è¯´å®è¯""çœŸçš„ç»äº†""ä¸€å¼€å§‹æˆ‘ä¹Ÿæ‹…å¿ƒ"ç­‰çœŸå®æ„Ÿå—
ğŸ”¥ çŸ­å¥+emojiï¼šæ¯å¥è¯ä¸è¶…è¿‡20å­—ï¼Œé€‚å½“åŠ ğŸ’¡âœ¨ğŸ”¥â—ç­‰emoji
ğŸ”¥ æœ‰èŠ‚å¥ï¼šç”¨"ï¼"æ¯”"ã€‚"å¤šï¼Œè¯»èµ·æ¥è¦æœ‰æ¿€åŠ¨æ„Ÿ
ğŸ”¥ ç»“å°¾è¦äº’åŠ¨ï¼š"å§å¦¹ä»¬å†²ï¼""æœ‰åŒæ¬¾å®å®çš„å¦ˆå¦ˆè¯„è®ºåŒºä¸¾æ‰‹ğŸ™‹â€â™€ï¸"

ã€å†…å®¹ç»“æ„ã€‘ï¼ˆæŒ‰è¿™ä¸ªé¡ºåºå†™ï¼Œè‡ªç„¶è¿‡æ¸¡ï¼‰
1. å¼€ç¯‡é’©å­ï¼šä½œä¸ºè‚²å©´å¸ˆ/è¥å…»å¸ˆï¼Œè¯´è¯´å¦ˆå¦ˆä»¬æœ€æ‹…å¿ƒçš„æˆ·å¤–å¸¦å¨ƒæ•æ•é—®é¢˜ï¼ˆçº¦70å­—ï¼‰
2. ç—›ç‚¹å…±é¸£ï¼šæˆ‘å›½åˆç”Ÿå®å®æ•æ•ç‡40%ï¼Œæœ‰å®¶æ—å²é£™åˆ°80%ï¼Œå¤ªå¯æ€•äº†ï¼ˆçº¦70å­—ï¼‰
3. ç§‘å­¦æ”¯æ‹›ï¼šç¬¬ä¸€å£å¥¶ç²‰é€‰å¯¹å¾ˆå…³é”®ï¼Œæ¨èé€‚åº¦æ°´è§£é…æ–¹ï¼ˆçº¦200å­—ï¼‰
4. äº§å“ç§è‰ï¼šé‡ç‚¹ä»‹ç»èƒ½æ©å…¨æŠ¤çš„æ°´è§£æŠ€æœ¯ã€è‡ªæŠ¤åŠ›é…æ–¹ã€è¥å…»æˆåˆ†ï¼ˆçº¦450å­—ï¼Œèå…¥10å¥è¯æœ¯ï¼‰
5. æ”¶å°¾å·å¬ï¼šæƒ³å¸¦å¨ƒæ”¾å¿ƒç©ï¼Œé€‰å¯¹å¥¶ç²‰æ˜¯ç¬¬ä¸€æ­¥ï¼ï¼ˆçº¦60å­—ï¼‰

ã€å…¶ä»–è¦æ±‚ã€‘
- æä¾›3ä¸ªæ ‡é¢˜å¤‡é€‰ï¼ˆå¿…é¡»å«ï¼šé€‚åº¦æ°´è§£ã€é˜²æ•ã€ç§‘æ™®ï¼‰
- æä¾›10ä¸ªä»¥ä¸Šè¯é¢˜æ ‡ç­¾ï¼ˆå¿…é¡»å«ï¼š#èƒ½æ©å…¨æŠ¤ #é€‚åº¦æ°´è§£ #é€‚åº¦æ°´è§£å¥¶ç²‰æ¨è #ç¬¬ä¸€å£å¥¶ç²‰ï¼‰
- å–ç‚¹é¡ºåºï¼šé˜²æ•æ°´è§£æŠ€æœ¯ â†’ è‡ªæŠ¤åŠ› â†’ åŸºç¡€è¥å…»
- ç¦è¯æ›¿æ¢ï¼šè¿‡æ•â†’æ•æ•ï¼Œé¢„é˜²â†’é˜²æ•ï¼Œæ–°ç”Ÿå„¿â†’åˆç”Ÿå®å®ï¼Œç”Ÿé•¿å‘è‚²â†’æˆé•¿
- ç»å¯¹ç¦æ­¢å‡ºç°ï¼šæ•å®ã€å¥¶ç“¶ã€å¥¶å˜´ã€ç–¾ç—…ã€æ²»ç–—ã€å…ç–«

ã€è¾“å‡ºæ ¼å¼ã€‘
### æ ‡é¢˜å¤‡é€‰ï¼ˆ3ä¸ªï¼‰
1. xxx
2. xxx
3. xxx

### æ­£æ–‡ï¼ˆ800-900å­—ï¼Œå¿…é¡»å†™å¤Ÿï¼ï¼‰
ï¼ˆè¿™é‡Œè¾“å‡ºå®Œæ•´æ­£æ–‡ï¼Œè¦æœ‰å°çº¢ä¹¦çˆ†æ–‡çš„æ´»äººæ„Ÿï¼ï¼‰

### è¯é¢˜æ ‡ç­¾
#èƒ½æ©å…¨æŠ¤ #é€‚åº¦æ°´è§£ ...ï¼ˆ10ä¸ªä»¥ä¸Šï¼‰

---
ã€éœ€è¦æ”¹å†™çš„KOLåŸç¨¿ã€‘
{content}"""

# ========== å·¥å…·å‡½æ•° ==========
def read_docx(file):
    doc = Document(io.BytesIO(file.read()))
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    return "\n".join(text)

def call_llm_api(prompt):
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return "Error: æœªè®¾ç½®OPENAI_API_KEYç¯å¢ƒå˜é‡"
    url = "https://api.openai.com/v1/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
    data = {"model": "gpt-4o", "max_tokens": 4000, "temperature": 0.8, "messages": [{"role": "user", "content": prompt}]}
    req = urllib.request.Request(url, data=json.dumps(data).encode('utf-8'), headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=120) as response:
            result = json.loads(response.read().decode('utf-8'))
            return result["choices"][0]["message"]["content"]
    except Exception as e:
        return f"Error: {str(e)}"

def count_chinese(text):
    return len(re.findall(r'[\u4e00-\u9fff]', text))

def extract_tags(content):
    return re.findall(r'#[\w\u4e00-\u9fff]+', content)

def extract_title(content):
    for line in content.split('\n'):
        line = line.strip()
        if line and not line.startswith('#'):
            return line
    return ""

def check_forbidden_word(content, word):
    """æ£€æŸ¥ç¦è¯æ˜¯å¦å‡ºç°ï¼Œè¿”å›è¿è§„ä½ç½®åˆ—è¡¨"""
    exceptions = FORBIDDEN_EXCEPTIONS.get(word, [])
    violations = []
    start = 0
    while True:
        idx = content.find(word, start)
        if idx == -1:
            break
        ctx = content[max(0, idx - 10):idx + len(word) + 10]
        is_exception = any(exc in ctx for exc in exceptions)
        if not is_exception:
            violations.append({"pos": idx, "context": ctx})
        start = idx + 1
    return violations

def auto_insert_fixed_phrases(content):
    """è‡ªåŠ¨æ’å…¥ç¼ºå¤±çš„ä¸å¯ä¿®æ”¹è¯æœ¯ï¼Œè¿”å›ä¿®å¤åçš„å†…å®¹"""
    # æ£€æŸ¥å“ªäº›è¯æœ¯ç¼ºå¤±
    missing_by_cat = {"é˜²æ•-æ°´è§£æŠ€æœ¯": [], "è‡ªæŠ¤åŠ›": [], "åŸºç¡€è¥å…»": []}
    for item in FIXED_SELLING_POINTS:
        if item["text"] not in content:
            missing_by_cat[item["category"]].append(item["text"])

    # å¦‚æœæ²¡æœ‰ç¼ºå¤±ï¼Œç›´æ¥è¿”å›
    total_missing = sum(len(v) for v in missing_by_cat.values())
    if total_missing == 0:
        return content, 0

    # æ‰¾åˆ°æ­£æ–‡éƒ¨åˆ†
    body_match = re.search(r'###\s*æ­£æ–‡[^#]*?\n(.*?)(?=###|$)', content, re.DOTALL)
    if not body_match:
        # å¦‚æœæ²¡æœ‰æ˜ç¡®çš„æ­£æ–‡æ ‡è®°ï¼Œå°è¯•åœ¨æ•´ä¸ªå†…å®¹ä¸­æ’å…¥
        body_match = re.search(r'([\s\S]+)', content)

    if not body_match:
        return content, 0

    body = body_match.group(1)
    modified_body = body
    inserted = 0

    # å®šä¹‰æ¯ä¸ªç±»åˆ«çš„é”šç‚¹å…³é”®è¯ï¼ˆç”¨äºæ‰¾åˆ°æ’å…¥ä½ç½®ï¼‰
    category_anchors = {
        "é˜²æ•-æ°´è§£æŠ€æœ¯": ["æ°´è§£", "é˜²æ•", "è›‹ç™½", "GINI", "è‡´æ•"],
        "è‡ªæŠ¤åŠ›": ["è‡ªæŠ¤", "HMO", "åŒèŒ", "ä¿æŠ¤åŠ›", "èŒèŒ", "è‚šè‚š"],
        "åŸºç¡€è¥å…»": ["è¥å…»", "ç»´ç”Ÿç´ ", "ä¹³ç³–", "å£å‘³"],
    }

    for cat, missing_phrases in missing_by_cat.items():
        if not missing_phrases:
            continue

        # æ‰¾åˆ°è¯¥ç±»åˆ«çš„é”šç‚¹ä½ç½®
        anchors = category_anchors.get(cat, [])
        best_pos = -1
        for anchor in anchors:
            pos = modified_body.find(anchor)
            if pos != -1:
                # æ‰¾åˆ°è¿™ä¸ªé”šç‚¹æ‰€åœ¨å¥å­çš„æœ«å°¾
                end_pos = modified_body.find("ã€‚", pos)
                if end_pos == -1:
                    end_pos = modified_body.find("ï¼", pos)
                if end_pos == -1:
                    end_pos = modified_body.find("\n", pos)
                if end_pos != -1:
                    best_pos = end_pos + 1
                    break

        # å¦‚æœæ‰¾ä¸åˆ°é”šç‚¹ï¼Œå°±æ’å…¥åˆ°æ­£æ–‡æœ«å°¾
        if best_pos == -1:
            best_pos = len(modified_body)

        # æ’å…¥ç¼ºå¤±çš„è¯æœ¯
        for phrase in missing_phrases:
            insert_text = phrase
            # ç¡®ä¿å‰åæœ‰é€‚å½“çš„æ ‡ç‚¹å’Œæ¢è¡Œ
            if best_pos > 0 and modified_body[best_pos-1] not in "ã€‚ï¼\n":
                insert_text = "ã€‚" + insert_text
            if not insert_text.endswith(("ã€‚", "ï¼")):
                insert_text += "ã€‚"

            modified_body = modified_body[:best_pos] + insert_text + modified_body[best_pos:]
            best_pos += len(insert_text)
            inserted += 1

    # æ›¿æ¢åŸå†…å®¹ä¸­çš„æ­£æ–‡éƒ¨åˆ†
    result = content.replace(body, modified_body)
    return result, inserted

def run_all_checks(content):
    """è¿è¡Œå…¨éƒ¨å®¡æ ¸æ£€æŸ¥ï¼Œè¿”å›ç»“æœå­—å…¸"""
    results = {}
    title = extract_title(content)
    tags = extract_tags(content)
    word_count = count_chinese(content)

    # å®¡æ ¸1: å–ç‚¹é¡ºåº
    positions = {}
    for cat, phrases in ORDER_ANCHORS.items():
        min_pos = float('inf')
        for p in phrases:
            idx = content.find(p)
            if idx != -1 and idx < min_pos:
                min_pos = idx
        positions[cat] = min_pos if min_pos != float('inf') else -1
    cats = ["é˜²æ•-æ°´è§£æŠ€æœ¯", "è‡ªæŠ¤åŠ›", "åŸºç¡€è¥å…»"]
    order_ok = True
    order_details = []
    for i, cat in enumerate(cats):
        pos = positions[cat]
        found = pos != -1
        order_details.append({"category": cat, "position": pos, "found": found})
        if not found:
            order_ok = False
        elif i > 0 and positions[cats[i - 1]] != -1 and pos < positions[cats[i - 1]]:
            order_ok = False
    results["check1"] = {"status": "pass" if order_ok else "fail", "details": order_details}

    # å®¡æ ¸2: å­—æ•°ï¼ˆ800-900å­—ï¼‰
    results["check2"] = {"status": "pass" if 800 <= word_count <= 900 else "fail", "count": word_count}

    # å®¡æ ¸3: æ ‡é¢˜æ•°é‡
    results["check3"] = {"status": "fail", "count": 1, "title": title,
                         "note": "KOLåˆç¨¿é€šå¸¸åªæœ‰1ä¸ªæ ‡é¢˜ï¼Œéœ€æä¾›3ä¸ªå¤‡é€‰"}

    # å®¡æ ¸4: æ ‡ç­¾
    missing_tags = [t for t in REQUIRED_TAGS if t not in tags]
    results["check4"] = {
        "status": "pass" if len(tags) >= 10 and not missing_tags else "fail",
        "count": len(tags), "missing": missing_tags, "tags": tags,
    }

    # å®¡æ ¸5: å…³é”®è¯
    kw_items = []
    for w in TITLE_KEYWORDS:
        kw_items.append({"scope": "æ ‡é¢˜", "word": w, "found": w in title})
    for w in BODY_KEYWORDS:
        kw_items.append({"scope": "æ­£æ–‡", "word": w, "found": w in content})
    for w in COVER_KEYWORDS:
        kw_items.append({"scope": "å°é¢(éœ€äººå·¥ç¡®è®¤)", "word": w, "found": w in content})
    results["check5"] = {
        "status": "pass" if all(r["found"] for r in kw_items) else "fail",
        "items": kw_items,
    }

    # å®¡æ ¸6: ç¦è¯
    fw_items = []
    for cat, words in FORBIDDEN_WORDS.items():
        for w in words:
            violations = check_forbidden_word(content, w)
            rep = FORBIDDEN_REPLACEMENTS.get(w, "åˆ é™¤")
            fw_items.append({
                "category": cat, "word": w,
                "found": len(violations) > 0,
                "violations": violations,
                "replacement": rep,
            })
    results["check6"] = {
        "status": "fail" if any(r["found"] for r in fw_items) else "pass",
        "items": fw_items,
    }

    # å®¡æ ¸7: å¿…æéœ€æ¶¦è‰²å–ç‚¹
    pp_items = []
    for sp in PARAPHRASE_SELLING_POINTS:
        found = sp["fragment"] in content
        pp_items.append({**sp, "found": found})
    results["check7"] = {
        "status": "pass" if all(r["found"] for r in pp_items) else "fail",
        "items": pp_items,
    }

    # å®¡æ ¸8: å¿…æä¸å¯ä¿®æ”¹å–ç‚¹
    fp_items = []
    for sp in FIXED_SELLING_POINTS:
        found = sp["text"] in content
        fp_items.append({**sp, "found": found})
    results["check8"] = {
        "status": "pass" if all(r["found"] for r in fp_items) else "fail",
        "items": fp_items,
    }

    # å®¡æ ¸9: å…è®¸åˆ å‡çš„å–ç‚¹
    op_items = []
    for sp in OPTIONAL_SELLING_POINTS:
        found = sp["fragment"] in content
        op_items.append({**sp, "found": found})
    results["check9"] = {"items": op_items}

    return results

def apply_adopted_changes(original, adopted_map, edit_map, check_results):
    """æ ¹æ®é‡‡çº³çš„ä¿®æ”¹å»ºè®®ç”Ÿæˆä¿®æ”¹åçš„æ–‡æœ¬"""
    modified = original
    changes = []

    # åº”ç”¨ç¦è¯æ›¿æ¢ (å®¡æ ¸6)
    if "check6" in check_results:
        for i, item in enumerate(check_results["check6"]["items"]):
            key = f"c6_{i}"
            if adopted_map.get(key) and item["found"]:
                old_word = item["word"]
                new_word = edit_map.get(key, item["replacement"])
                if old_word in modified:
                    # å°Šé‡ä¾‹å¤–
                    exceptions = FORBIDDEN_EXCEPTIONS.get(old_word, [])
                    if exceptions:
                        # é€ä¸ªä½ç½®æ›¿æ¢ï¼Œè·³è¿‡ä¾‹å¤–
                        result = []
                        start = 0
                        while True:
                            idx = modified.find(old_word, start)
                            if idx == -1:
                                result.append(modified[start:])
                                break
                            ctx = modified[max(0, idx - 10):idx + len(old_word) + 10]
                            is_exc = any(exc in ctx for exc in exceptions)
                            if is_exc:
                                result.append(modified[start:idx + len(old_word)])
                            else:
                                result.append(modified[start:idx])
                                result.append(new_word)
                                changes.append({"old": old_word, "new": new_word})
                            start = idx + len(old_word)
                        modified = "".join(result)
                    else:
                        modified = modified.replace(old_word, new_word)
                        changes.append({"old": old_word, "new": new_word})

    # è¡¥å……ç¼ºå¤±æ ‡ç­¾ (å®¡æ ¸4)
    if "check4" in check_results:
        missing = check_results["check4"].get("missing", [])
        for i, tag in enumerate(missing):
            key = f"c4_{i}"
            if adopted_map.get(key):
                if tag not in modified:
                    modified = modified.rstrip() + " " + tag
                    changes.append({"old": "", "new": tag})

    return modified, changes

def highlight_diff(text, changes, mode="original"):
    """å¯¹æ–‡æœ¬ä¸­çš„ä¿®æ”¹éƒ¨åˆ†è¿›è¡Œé«˜äº®"""
    html = text
    for c in changes:
        if mode == "original" and c["old"]:
            html = html.replace(
                c["old"],
                f'<span style="background:#c8e6c9;padding:1px 4px;border-radius:3px;font-weight:bold;">{c["old"]}</span>'
            )
        elif mode == "modified" and c["new"]:
            html = html.replace(
                c["new"],
                f'<span style="background:#f8bbd0;padding:1px 4px;border-radius:3px;font-weight:bold;">{c["new"]}</span>'
            )
    return html.replace('\n', '<br>')

# ========== é¡µé¢é…ç½® ==========
st.set_page_config(page_title="èµæ„AIå®¡ç¨¿ç³»ç»Ÿ", page_icon="ğŸ¤–", layout="wide")

st.markdown("""
<style>
.block-container {padding-top: 1rem !important; padding-bottom: 1rem !important;}
/* å¯¼èˆªæ  */
.nav-bar {
    display: flex; gap: 0; margin-bottom: 20px; border-radius: 10px; overflow: hidden;
    border: 2px solid #ddd;
}
.nav-item {
    flex: 1; text-align: center; padding: 12px 10px; font-weight: bold; font-size: 15px;
    cursor: default;
}
.nav-part1 { background: #eef1fa; color: #2c3e6b; border-right: 2px solid #ddd; }
.nav-part2 { background: #f0faf4; color: #2e7d32; border-right: 2px solid #ddd; }
.nav-part3 { background: #fff8e1; color: #f57c00; border-right: 2px solid #ddd; }
.nav-part4 { background: #fce4ec; color: #c2185b; }
/* é»˜è®¤å®¹å™¨æ ·å¼ - é˜²æ­¢æ ·å¼æ³„æ¼ */
[data-testid="stVerticalBlockBorderWrapper"] {
    background-color: white !important;
}
/* Part1 è“ç´«è‰² */
[data-testid="stVerticalBlockBorderWrapper"]:has(#part1-marker) {
    background-color: #f5f3ff !important;
    border: 2px solid #c4b5fd !important;
    border-radius: 12px !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part1-marker) button[kind="primary"] {
    background-color: #7c3aed !important; border-color: #7c3aed !important; color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part1-marker) button[kind="primary"]:hover {
    background-color: #6d28d9 !important; border-color: #6d28d9 !important;
}
/* Part2 ç»¿è‰² */
[data-testid="stVerticalBlockBorderWrapper"]:has(#part2-marker) {
    background-color: #edf7f0 !important;
    border: 2px solid #b4dfc6 !important;
    border-radius: 12px !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part2-marker) button[kind="primary"] {
    background-color: #4caf50 !important; border-color: #4caf50 !important; color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part2-marker) button[kind="primary"]:hover {
    background-color: #388e3c !important; border-color: #388e3c !important;
}
/* Part3 æ©™è‰² */
[data-testid="stVerticalBlockBorderWrapper"]:has(#part3-marker) {
    background-color: #fff8e1 !important;
    border: 2px solid #ffcc80 !important;
    border-radius: 12px !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part3-marker) button[kind="primary"] {
    background-color: #ff9800 !important; border-color: #ff9800 !important; color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part3-marker) button[kind="primary"]:hover {
    background-color: #f57c00 !important; border-color: #f57c00 !important;
}
/* Part4 ç²‰è‰² */
[data-testid="stVerticalBlockBorderWrapper"]:has(#part4-marker) {
    background-color: #fce4ec !important;
    border: 2px solid #f48fb1 !important;
    border-radius: 12px !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part4-marker) button[kind="primary"] {
    background-color: #e91e63 !important; border-color: #e91e63 !important; color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part4-marker) button[kind="primary"]:hover {
    background-color: #c2185b !important; border-color: #c2185b !important;
}
/* å®¡æ ¸å¡ç‰‡ */
.check-header-pass {
    background: #e8f5e9; border-left: 5px solid #4caf50; padding: 10px 15px;
    margin: 10px 0 5px 0; border-radius: 0 8px 8px 0; font-weight: bold; font-size: 15px;
}
.check-header-fail {
    background: #fce4ec; border-left: 5px solid #e57373; padding: 10px 15px;
    margin: 10px 0 5px 0; border-radius: 0 8px 8px 0; font-weight: bold; font-size: 15px;
}
.check-header-info {
    background: #fff8e1; border-left: 5px solid #ffc107; padding: 10px 15px;
    margin: 10px 0 5px 0; border-radius: 0 8px 8px 0; font-weight: bold; font-size: 15px;
}
/* æ–‡ä»¶ä¸Šä¼ ä¸­æ–‡åŒ– */
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] p {font-size: 0 !important;}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] p::after {content: "å°†æ–‡ä»¶æ‹–åˆ°æ­¤å¤„ä¸Šä¼ "; font-size: 14px !important;}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] button {font-size: 0 !important; position: relative;}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] button::after {content: "é€‰æ‹©æ–‡ä»¶"; font-size: 14px !important; position: absolute;}
/* æ‰€æœ‰ä¸‹è½½æŒ‰é’®ç»Ÿä¸€ç´«è‰² */
[data-testid="stDownloadButton"] > button {
    background-color: #7c3aed !important;
    border-color: #7c3aed !important;
    color: white !important;
}
[data-testid="stDownloadButton"] > button:hover {
    background-color: #6d28d9 !important;
    border-color: #6d28d9 !important;
}
</style>
""", unsafe_allow_html=True)

# ========== æ ‡é¢˜ ==========
st.markdown("""
<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 12px; padding: 20px 25px; margin-bottom: 15px;">
    <h2 style="color: white; margin: 0;">ğŸ¤– èµæ„AI Â· å°çº¢ä¹¦KOLå®¡ç¨¿ç³»ç»Ÿ for å…”å­ğŸ°</h2>
</div>
""", unsafe_allow_html=True)

# å¯¼èˆªæ 
st.markdown("""
<div class="nav-bar">
    <div class="nav-item nav-part1">Part 1 Â· å…«å¤§å®¡æ ¸</div>
    <div class="nav-item nav-part2">Part 2 Â· äººè¯ä¿®æ”¹</div>
    <div class="nav-item nav-part3">Part 3 Â· å¤æ ¸æ£€æŸ¥</div>
    <div class="nav-item nav-part4">Part 4 Â· ç»ˆç¨¿å®Œæˆ</div>
</div>
""", unsafe_allow_html=True)

# ========== Session State åˆå§‹åŒ– ==========
if 'kol_content' not in st.session_state:
    st.session_state.kol_content = ""
if 'audit_results' not in st.session_state:
    st.session_state.audit_results = None
if 'audit_adopted' not in st.session_state:
    st.session_state.audit_adopted = {}
if 'audit_edits' not in st.session_state:
    st.session_state.audit_edits = {}
if 'modified_content' not in st.session_state:
    st.session_state.modified_content = ""
if 'diff_changes' not in st.session_state:
    st.session_state.diff_changes = []
if 'renhua_result' not in st.session_state:
    st.session_state.renhua_result = ""
if 'renhua_adopted' not in st.session_state:
    st.session_state.renhua_adopted = False
if 'recheck_content' not in st.session_state:
    st.session_state.recheck_content = ""
if 'recheck_results' not in st.session_state:
    st.session_state.recheck_results = None
if 'final_content' not in st.session_state:
    st.session_state.final_content = ""
if 'final_ready' not in st.session_state:
    st.session_state.final_ready = False

# ========== ç¨¿ä»¶æ–¹å‘é€‰æ‹© ==========
DIRECTION_OPTIONS = [
    "è¯·é€‰æ‹©ç¨¿ä»¶æ–¹å‘...",
    "æ–¹å‘1.ã€è‚²å©´å¸ˆé˜²æ•ç§‘æ™®ã€‘",
    "æ–¹å‘2.ã€å•å“åˆ†äº«ã€‘",
    "æ–¹å‘3.ã€åå‘ç»éªŒåˆ†äº«-å®¶æ—è¿‡æ•å²ã€‘",
    "æ–¹å‘4.ã€åå‘ç»éªŒåˆ†äº«-å‰–è…¹äº§ã€‘",
    "æ–¹å‘5.ã€é˜²æ•å¾…äº§åŒ…åˆ†äº«-å­•æ™šæ•æ„Ÿã€‘",
    "æ–¹å‘6.ã€é˜²æ•å¾…äº§åŒ…åˆ†äº«-å‰–è…¹äº§ã€‘",
    "æ–¹å‘7.ã€å…»å® å®¶åº­ã€‘",
    "æ–¹å‘8.ã€a2VSç¬¬ä¸€å£é¡¶é…ã€‘",
    "æ–¹å‘9.ã€èƒ½æ©å…¨æŠ¤è´µæœ‰æ‰€å€¼ã€‘",
    "æ–¹å‘10.ã€èƒ½æ©å…¨æŠ¤+è¶…å¯èƒ½æ©å®¶æ—æµ‹è¯„ã€‘",
    "æ–¹å‘11.ã€é˜²æ•ç«å“æµ‹è¯„ã€‘",
    "æ–¹å‘12.ã€è·¨å¢ƒèƒ½æ©å…¨æµ‹è¯„ã€‘",
]

if 'selected_direction' not in st.session_state:
    st.session_state.selected_direction = DIRECTION_OPTIONS[0]

# ========== è¾“å…¥åŒº ==========
dir_col, date_col = st.columns([3, 1])
with dir_col:
    selected_dir = st.selectbox("æœ¬ç¨¿ä»¶ç¬¦åˆæ–¹å‘", DIRECTION_OPTIONS, key="direction_select")
    st.session_state.selected_direction = selected_dir
with date_col:
    st.caption(f"å½“å‰æ—¥æœŸ: {TODAY}")

upload_col, paste_col = st.columns(2)
with upload_col:
    kol_file = st.file_uploader("ä¸Šä¼ KOLç¨¿ä»¶ (.docx)", type=["docx"], key="kol_file")
with paste_col:
    kol_text = st.text_area("æˆ–ç²˜è´´ç¨¿ä»¶å†…å®¹", height=120, placeholder="åœ¨æ­¤ç²˜è´´KOLç¨¿ä»¶å†…å®¹...", key="kol_text")

if kol_file:
    kol_file.seek(0)
    st.session_state.kol_content = read_docx(kol_file)
elif kol_text:
    st.session_state.kol_content = kol_text

# ================================================================
# Part 1: å…«å¤§å®¡æ ¸
# ================================================================
with st.container(border=True):
    st.markdown('<div id="part1-marker"></div>', unsafe_allow_html=True)
    st.markdown("#### Part 1 Â· å…«å¤§å®¡æ ¸")
    st.caption("æœ¬åœ°Pythoné€é¡¹æ£€æŸ¥ï¼Œå…¨éƒ¨ç»“æœä»¥è¡¨æ ¼å±•ç¤ºï¼Œå‘ç°é—®é¢˜å¯ç¼–è¾‘å»ºè®®å¹¶é‡‡çº³ä¿å­˜")

    if not st.session_state.kol_content:
        st.info("è¯·å…ˆä¸Šä¼ æˆ–ç²˜è´´KOLç¨¿ä»¶")
    else:
        # å¼€å§‹å®¡æ ¸æŒ‰é’®
        if st.button("å¼€å§‹å…«å¤§å®¡æ ¸", key="btn_audit", use_container_width=True, type="primary"):
            st.session_state.audit_results = run_all_checks(st.session_state.kol_content)
            st.session_state.audit_adopted = {}
            st.session_state.audit_edits = {}
            st.session_state.modified_content = ""
            st.session_state.diff_changes = []
            st.rerun()

        # æ˜¾ç¤ºå®¡æ ¸ç»“æœ
        if st.session_state.audit_results:
            r = st.session_state.audit_results
            content = st.session_state.kol_content

            # --- ç»Ÿè®¡æ¦‚è§ˆ ---
            pass_count = sum(1 for k in ["check1","check2","check3","check4","check5","check6","check7","check8"]
                           if r.get(k, {}).get("status") == "pass")
            fail_count = 8 - pass_count
            m1, m2, m3 = st.columns(3)
            m1.metric("é€šè¿‡", f"{pass_count}/8")
            m2.metric("éœ€ä¿®æ”¹", f"{fail_count}")
            m3.metric("ç¨¿ä»¶å­—æ•°", f"{r['check2']['count']}")

            # ==========================================
            # å®¡æ ¸1: å–ç‚¹é¡ºåº
            # ==========================================
            s1 = r["check1"]["status"]
            icon1 = "âœ…" if s1 == "pass" else "âŒ"
            cls1 = "check-header-pass" if s1 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls1}">{icon1} å®¡æ ¸1ï¼šå–ç‚¹é¡ºåºï¼ˆé˜²æ•-æ°´è§£æŠ€æœ¯ â†’ è‡ªæŠ¤åŠ› â†’ åŸºç¡€è¥å…»ï¼‰</div>', unsafe_allow_html=True)
            rows1 = ""
            for d in r["check1"]["details"]:
                found = "âœ… å·²å‡ºç°" if d["found"] else "âŒ æœªå‡ºç°"
                pos = f"ä½ç½®: {d['position']}" if d["found"] else "â€”"
                bg = "#f0fff4" if d["found"] else "#fff5f5"
                rows1 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;">{d["category"]}</td><td style="border:1px solid #ddd;padding:6px 8px;">{found}</td><td style="border:1px solid #ddd;padding:6px 8px;">{pos}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">å–ç‚¹ç±»åˆ«</th><th style="border:1px solid #ddd;padding:8px;">æ£€æŸ¥ç»“æœ</th><th style="border:1px solid #ddd;padding:8px;">ä½ç½®</th></tr></thead>
            <tbody>{rows1}</tbody></table>''', unsafe_allow_html=True)
            if s1 == "fail":
                st.text_input("ä¿®æ”¹å»ºè®®", value="è¯·è°ƒæ•´æ®µè½é¡ºåºï¼šå…ˆå†™é˜²æ•-æ°´è§£æŠ€æœ¯ï¼Œå†å†™è‡ªæŠ¤åŠ›ï¼Œæœ€åå†™åŸºç¡€è¥å…»", key="edit_c1", disabled=False)
                st.checkbox("é‡‡çº³", key="adopt_c1", value=True)

            # ==========================================
            # å®¡æ ¸2: å­—æ•°æ£€æŸ¥
            # ==========================================
            s2 = r["check2"]["status"]
            icon2 = "âœ…" if s2 == "pass" else "âŒ"
            cls2 = "check-header-pass" if s2 == "pass" else "check-header-fail"
            wc = r["check2"]["count"]
            st.markdown(f'<div class="{cls2}">{icon2} å®¡æ ¸2ï¼šå­—æ•°æ£€æŸ¥ï¼ˆ{wc}å­—ï¼Œè¦æ±‚800-900å­—ï¼‰</div>', unsafe_allow_html=True)
            if s2 == "fail":
                st.warning(f"è¶…å‡º {wc - 900} å­—ï¼Œè¯·ç²¾ç®€å†…å®¹")
                wc_hint = "éœ€æ‰©å……å†…å®¹" if wc < 800 else "éœ€ç²¾ç®€å†…å®¹"
                st.text_input("ä¿®æ”¹å»ºè®®", value=f"å½“å‰{wc}å­—ï¼Œ{wc_hint}è‡³800-900å­—", key="edit_c2")
                st.checkbox("é‡‡çº³", key="adopt_c2", value=True)

            # ==========================================
            # å®¡æ ¸3: æ ‡é¢˜æ•°é‡
            # ==========================================
            st.markdown(f'<div class="check-header-fail">âŒ å®¡æ ¸3ï¼šæ ‡é¢˜æ•°é‡ï¼ˆéœ€3ä¸ªå¤‡é€‰æ ‡é¢˜ï¼‰</div>', unsafe_allow_html=True)
            title = r["check3"]["title"]
            st.markdown(f'<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;"><thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">å½“å‰æ ‡é¢˜</th><th style="border:1px solid #ddd;padding:8px;">è¦æ±‚</th><th style="border:1px solid #ddd;padding:8px;">ç»“æœ</th></tr></thead><tbody><tr style="background:#fff5f5;"><td style="border:1px solid #ddd;padding:6px 8px;">{title[:50]}...</td><td style="border:1px solid #ddd;padding:6px 8px;">3ä¸ªå¤‡é€‰æ ‡é¢˜</td><td style="border:1px solid #ddd;padding:6px 8px;">âŒ ä»…1ä¸ª</td></tr></tbody></table>', unsafe_allow_html=True)
            st.caption("å»ºè®®ï¼šäººè¯ä¿®æ”¹é˜¶æ®µAIå°†è‡ªåŠ¨ç”Ÿæˆ3ä¸ªå¤‡é€‰æ ‡é¢˜")

            # ==========================================
            # å®¡æ ¸4: è¯é¢˜æ ‡ç­¾
            # ==========================================
            s4 = r["check4"]["status"]
            icon4 = "âœ…" if s4 == "pass" else "âŒ"
            cls4 = "check-header-pass" if s4 == "pass" else "check-header-fail"
            tc = r["check4"]["count"]
            st.markdown(f'<div class="{cls4}">{icon4} å®¡æ ¸4ï¼šè¯é¢˜æ ‡ç­¾ï¼ˆå½“å‰{tc}ä¸ªï¼Œè¦æ±‚10ä¸ªä»¥ä¸Šï¼‰</div>', unsafe_allow_html=True)

            rows4 = ""
            for tag in REQUIRED_TAGS:
                found = tag in r["check4"]["tags"]
                icon = "âœ…" if found else "âŒ ç¼ºå¤±"
                bg = "#f0fff4" if found else "#fff5f5"
                rows4 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;">{tag}</td><td style="border:1px solid #ddd;padding:6px 8px;">{icon}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">å¿…å«æ ‡ç­¾</th><th style="border:1px solid #ddd;padding:8px;">ç»“æœ</th></tr></thead>
            <tbody>{rows4}</tbody></table>''', unsafe_allow_html=True)

            if r["check4"]["missing"]:
                for mi, mtag in enumerate(r["check4"]["missing"]):
                    c4_col1, c4_col2 = st.columns([3, 1])
                    with c4_col1:
                        st.text_input(f"è¡¥å……æ ‡ç­¾", value=mtag, key=f"edit_c4_{mi}")
                    with c4_col2:
                        st.checkbox("é‡‡çº³", key=f"adopt_c4_{mi}", value=True)

            # ==========================================
            # å®¡æ ¸5: å…³é”®è¯
            # ==========================================
            s5 = r["check5"]["status"]
            icon5 = "âœ…" if s5 == "pass" else "âŒ"
            cls5 = "check-header-pass" if s5 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls5}">{icon5} å®¡æ ¸5ï¼šå¿…é¡»å‡ºç°å…³é”®è¯</div>', unsafe_allow_html=True)

            rows5 = ""
            for item in r["check5"]["items"]:
                found = "âœ… å·²åŒ…å«" if item["found"] else "âŒ ç¼ºå¤±"
                bg = "#f0fff4" if item["found"] else "#fff5f5"
                rows5 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;">{item["scope"]}</td><td style="border:1px solid #ddd;padding:6px 8px;font-weight:bold;">{item["word"]}</td><td style="border:1px solid #ddd;padding:6px 8px;">{found}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">æ£€æŸ¥èŒƒå›´</th><th style="border:1px solid #ddd;padding:8px;">å…³é”®è¯</th><th style="border:1px solid #ddd;padding:8px;">ç»“æœ</th></tr></thead>
            <tbody>{rows5}</tbody></table>''', unsafe_allow_html=True)

            missing_kw = [item for item in r["check5"]["items"] if not item["found"]]
            for ki, kw_item in enumerate(missing_kw):
                c5_col1, c5_col2 = st.columns([3, 1])
                with c5_col1:
                    st.text_input(f"ä¿®æ”¹å»ºè®®", value=f"è¯·åœ¨{kw_item['scope']}ä¸­åŠ å…¥ã€Œ{kw_item['word']}ã€", key=f"edit_c5_{ki}")
                with c5_col2:
                    st.checkbox("é‡‡çº³", key=f"adopt_c5_{ki}", value=True)

            # ==========================================
            # å®¡æ ¸6: ç¦è¯/ç¦ç”¨è¡¨è¾¾
            # ==========================================
            s6 = r["check6"]["status"]
            icon6 = "âœ…" if s6 == "pass" else "âŒ"
            cls6 = "check-header-pass" if s6 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls6}">{icon6} å®¡æ ¸6ï¼šç¦è¯/ç¦ç”¨è¡¨è¾¾æ£€æŸ¥</div>', unsafe_allow_html=True)

            rows6 = ""
            for item in r["check6"]["items"]:
                if item["found"]:
                    icon = "âŒ å‡ºç°äº†"
                    bg = "#fff5f5"
                    ctx_list = item.get("violations", [])
                    ctx_str = "ã€".join([f'"{v["context"].strip()}"' for v in ctx_list[:2]])
                else:
                    icon = "âœ… æœªå‡ºç°"
                    bg = "#f0fff4"
                    ctx_str = "â€”"
                rows6 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;">{item["category"]}</td><td style="border:1px solid #ddd;padding:6px 8px;font-weight:bold;">{item["word"]}</td><td style="border:1px solid #ddd;padding:6px 8px;">{icon}</td><td style="border:1px solid #ddd;padding:6px 8px;font-size:12px;">{ctx_str}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">ç±»å‹</th><th style="border:1px solid #ddd;padding:8px;">ç¦è¯</th><th style="border:1px solid #ddd;padding:8px;">ç»“æœ</th><th style="border:1px solid #ddd;padding:8px;">ä¸Šä¸‹æ–‡</th></tr></thead>
            <tbody>{rows6}</tbody></table>''', unsafe_allow_html=True)

            found_forbidden = [item for item in r["check6"]["items"] if item["found"]]
            for fi, fw in enumerate(found_forbidden):
                c6_col1, c6_col2 = st.columns([3, 1])
                with c6_col1:
                    default_rep = fw["replacement"]
                    edited = st.text_input(f"ã€Œ{fw['word']}ã€æ›¿æ¢ä¸º", value=default_rep, key=f"edit_c6_{fi}")
                    st.session_state.audit_edits[f"c6_{r['check6']['items'].index(fw)}"] = edited
                with c6_col2:
                    st.checkbox("é‡‡çº³", key=f"adopt_c6_{fi}", value=True)

            # ==========================================
            # å®¡æ ¸7: å¿…æéœ€æ¶¦è‰²å–ç‚¹ï¼ˆä¸Šä¸‹è¡Œç»“æ„ï¼‰
            # ==========================================
            s7 = r["check7"]["status"]
            icon7 = "âœ…" if s7 == "pass" else "âŒ"
            cls7 = "check-header-pass" if s7 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls7}">{icon7} å®¡æ ¸7ï¼šå¿…æéœ€æ¶¦è‰²å–ç‚¹ï¼ˆ4å¤§æ–¹å‘ Â· 10å°æ–¹å‘ï¼‰</div>', unsafe_allow_html=True)

            # ä¸Šä¸‹è¡Œç»“æ„ï¼šæ¯ä¸ªå–ç‚¹ä¸€ä¸ªå¡ç‰‡ï¼Œå†…å®¹+å»ºè®®åœ¨ä¸€èµ·
            current_cat7 = ""
            pi_counter = 0
            for item in r["check7"]["items"]:
                # åˆ†ç±»æ ‡é¢˜
                if item["category"] != current_cat7:
                    current_cat7 = item["category"]
                    st.markdown(f'<div style="background:#e8eaf6;padding:6px 12px;margin-top:10px;border-radius:5px;font-weight:bold;color:#3949ab;">ğŸ“‚ å¤§æ–¹å‘ï¼š{current_cat7}</div>', unsafe_allow_html=True)

                found = item["found"]
                icon = "âœ…" if found else "âŒ"
                bg = "#f0fff4" if found else "#fff5f5"
                border_color = "#4caf50" if found else "#ef5350"

                # å¡ç‰‡ï¼šè¯æœ¯å†…å®¹
                st.markdown(f'''<div style="background:{bg};border-left:4px solid {border_color};padding:10px 15px;margin:6px 0;border-radius:0 8px 8px 0;">
                <div style="font-size:13px;"><b>å°æ–¹å‘{item["idx"]}</b> {icon}</div>
                <div style="font-size:13px;color:#333;margin-top:4px;line-height:1.6;">{item["text"]}</div>
                </div>''', unsafe_allow_html=True)

                # å¦‚æœæœªæ‰¾åˆ°ï¼Œæ˜¾ç¤ºå»ºè®®ç¼–è¾‘æ¡†ï¼ˆç´§è·Ÿåœ¨å†…å®¹ä¸‹æ–¹ï¼‰
                if not found:
                    c7_col1, c7_col2 = st.columns([4, 1])
                    with c7_col1:
                        st.text_input("ä¿®æ”¹å»ºè®®", value=f"éœ€æ¶¦è‰²åŠ å…¥ï¼š{item['text']}", key=f"edit_c7_{pi_counter}", label_visibility="collapsed")
                    with c7_col2:
                        st.checkbox("é‡‡çº³", key=f"adopt_c7_{pi_counter}", value=True)
                    pi_counter += 1

            # ==========================================
            # å®¡æ ¸8: å¿…æä¸å¯ä¿®æ”¹å–ç‚¹ï¼ˆä¸Šä¸‹è¡Œç»“æ„ï¼‰
            # ==========================================
            s8 = r["check8"]["status"]
            icon8 = "âœ…" if s8 == "pass" else "âŒ"
            cls8 = "check-header-pass" if s8 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls8}">{icon8} å®¡æ ¸8ï¼šå¿…æä¸å¯ä¿®æ”¹å–ç‚¹ï¼ˆ3å¤§åˆ‡è§’ Â· 10å°åˆ‡è§’ï¼Œå¿…é¡»å­—å­—ä¸å·®ï¼‰</div>', unsafe_allow_html=True)

            # ä¸Šä¸‹è¡Œç»“æ„ï¼šæ¯ä¸ªå–ç‚¹ä¸€ä¸ªå¡ç‰‡ï¼Œå†…å®¹+å»ºè®®åœ¨ä¸€èµ·
            current_cat8 = ""
            fpi_counter = 0
            for item in r["check8"]["items"]:
                # åˆ†ç±»æ ‡é¢˜
                if item["category"] != current_cat8:
                    current_cat8 = item["category"]
                    st.markdown(f'<div style="background:#fce4ec;padding:6px 12px;margin-top:10px;border-radius:5px;font-weight:bold;color:#c2185b;">ğŸ“‚ å¤§åˆ‡è§’ï¼š{current_cat8}</div>', unsafe_allow_html=True)

                found = item["found"]
                icon = "âœ…" if found else "âŒ"
                bg = "#f0fff4" if found else "#fff5f5"
                border_color = "#4caf50" if found else "#ef5350"

                # å¡ç‰‡ï¼šè¯æœ¯å†…å®¹ï¼ˆå¼ºè°ƒä¸å¯ä¿®æ”¹ï¼‰
                if found:
                    # å·²æ‰¾åˆ°ï¼šæ˜¾ç¤ºæ­£å¸¸å¡ç‰‡
                    st.markdown(f'''<div style="background:{bg};border-left:4px solid {border_color};padding:10px 15px;margin:6px 0;border-radius:0 8px 8px 0;">
                    <div style="font-size:13px;"><b>å°åˆ‡è§’{item["idx"]}</b> {icon} <span style="color:#4caf50;font-size:11px;">ï¼ˆå·²åŒ…å«ï¼‰</span></div>
                    <div style="font-size:13px;color:#333;margin-top:4px;line-height:1.6;font-weight:500;">{item["text"]}</div>
                    </div>''', unsafe_allow_html=True)
                else:
                    # æœªæ‰¾åˆ°ï¼šæ˜¾ç¤º"æ²¡æœ‰æåˆ°ï¼Œå»ºè®®å¢åŠ "
                    st.markdown(f'''<div style="background:{bg};border-left:4px solid {border_color};padding:10px 15px;margin:6px 0;border-radius:0 8px 8px 0;">
                    <div style="font-size:13px;"><b>å°åˆ‡è§’{item["idx"]}</b> {icon} <span style="color:#c62828;font-size:11px;font-weight:bold;">æ²¡æœ‰æåˆ°</span></div>
                    <div style="font-size:13px;color:#c62828;margin-top:4px;line-height:1.6;font-weight:600;">å»ºè®®å¢åŠ ï¼š<span style="color:#333;">{item["text"]}</span></div>
                    </div>''', unsafe_allow_html=True)

                # å¦‚æœæœªæ‰¾åˆ°ï¼Œæ˜¾ç¤ºå»ºè®®ç¼–è¾‘æ¡†ï¼ˆç´§è·Ÿåœ¨å†…å®¹ä¸‹æ–¹ï¼‰
                if not found:
                    c8_col1, c8_col2 = st.columns([4, 1])
                    with c8_col1:
                        st.text_input("ä¿®æ”¹å»ºè®®", value=f"å¿…é¡»åŸå°ä¸åŠ¨åŠ å…¥ï¼š{item['text']}", key=f"edit_c8_{fpi_counter}", label_visibility="collapsed")
                    with c8_col2:
                        st.checkbox("é‡‡çº³", key=f"adopt_c8_{fpi_counter}", value=True)
                    fpi_counter += 1

            # ==========================================
            # å®¡æ ¸9: å…è®¸åˆ å‡çš„å–ç‚¹
            # ==========================================
            st.markdown(f'<div class="check-header-info">â„¹ï¸ å®¡æ ¸9ï¼šå…è®¸åˆ å‡çš„å–ç‚¹ï¼ˆä»…ä¾›å‚è€ƒï¼‰</div>', unsafe_allow_html=True)
            rows9 = ""
            for item in r["check9"]["items"]:
                found = "âœ… å·²å‡ºç°" if item["found"] else "â€” æœªå‡ºç°ï¼ˆå¯åˆ å‡ï¼‰"
                bg = "#f0fff4" if item["found"] else "#fffde7"
                rows9 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;"><b>{item["category"]}</b></td><td style="border:1px solid #ddd;padding:6px 8px;font-size:12px;">{item["text"]}</td><td style="border:1px solid #ddd;padding:6px 8px;">{found}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">ç±»åˆ«</th><th style="border:1px solid #ddd;padding:8px;">å–ç‚¹å†…å®¹</th><th style="border:1px solid #ddd;padding:8px;">çŠ¶æ€</th></tr></thead>
            <tbody>{rows9}</tbody></table>''', unsafe_allow_html=True)

            # ==========================================
            # æ ‡å‡†å–ç‚¹ç¤ºä¾‹
            # ==========================================
            with st.expander("ğŸ“– æ ‡å‡†å–ç‚¹ç¤ºä¾‹ï¼ˆå‚è€ƒï¼‰", expanded=False):
                st.markdown(SELLING_POINT_EXAMPLE)

            # ==========================================
            # ä¿å­˜é‡‡çº³ + å·¦å³å¯¹æ¯”
            # ==========================================
            st.markdown("---")
            if st.button("ä¿å­˜æ‰€æœ‰é‡‡çº³ä¿®æ”¹ â†’ ç”Ÿæˆå¯¹æ¯”é¢„è§ˆ", key="btn_save_audit", use_container_width=True, type="primary"):
                # æ”¶é›†æ‰€æœ‰é‡‡çº³çŠ¶æ€
                adopted = {}
                edits = {}

                # å®¡æ ¸6 ç¦è¯æ›¿æ¢
                found_fw = [item for item in r["check6"]["items"] if item["found"]]
                for fi, fw in enumerate(found_fw):
                    real_idx = r["check6"]["items"].index(fw)
                    adopted[f"c6_{real_idx}"] = st.session_state.get(f"adopt_c6_{fi}", False)
                    edits[f"c6_{real_idx}"] = st.session_state.get(f"edit_c6_{fi}", fw["replacement"])

                # å®¡æ ¸4 ç¼ºå¤±æ ‡ç­¾
                missing_tags = r["check4"].get("missing", [])
                for mi, _ in enumerate(missing_tags):
                    adopted[f"c4_{mi}"] = st.session_state.get(f"adopt_c4_{mi}", False)

                st.session_state.audit_adopted = adopted
                st.session_state.audit_edits = edits

                modified, changes = apply_adopted_changes(
                    st.session_state.kol_content, adopted, edits, r
                )
                st.session_state.modified_content = modified
                st.session_state.diff_changes = changes
                st.rerun()

            # æ˜¾ç¤ºå¯¹æ¯”é¢„è§ˆ
            if st.session_state.modified_content:
                st.markdown("---")
                st.markdown("### å¯¹æ¯”é¢„è§ˆï¼ˆåŸæ–‡ vs ä¿®æ”¹åï¼‰")
                st.caption("ğŸŸ¢ ç»¿è‰² = åŸæ–‡ä¸­è¢«ä¿®æ”¹çš„éƒ¨åˆ† | ğŸ©· ç²‰è‰² = ä¿®æ”¹åçš„å†…å®¹")

                cmp_left, cmp_right = st.columns(2)
                with cmp_left:
                    st.markdown("**åŸæ–‡ï¼ˆç»¿è‰²æ ‡æ³¨ä¿®æ”¹å¤„ï¼‰**")
                    orig_html = highlight_diff(st.session_state.kol_content, st.session_state.diff_changes, "original")
                    st.markdown(f'<div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:15px;font-size:14px;line-height:2.0;">{orig_html}</div>', unsafe_allow_html=True)

                with cmp_right:
                    st.markdown("**ä¿®æ”¹åï¼ˆç²‰è‰²æ ‡æ³¨ä¿®æ”¹å¤„ï¼‰**")
                    mod_html = highlight_diff(st.session_state.modified_content, st.session_state.diff_changes, "modified")
                    st.markdown(f'<div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:15px;font-size:14px;line-height:2.0;">{mod_html}</div>', unsafe_allow_html=True)

                # ç¼–è¾‘ + é‡‡ç”¨
                with st.expander("éœ€è¦å¾®è°ƒï¼Ÿç‚¹å‡»ç¼–è¾‘ä¿®æ”¹åå†…å®¹", expanded=False):
                    edited_mod = st.text_area("ä¿®æ”¹åå†…å®¹ï¼ˆå¯ç¼–è¾‘ï¼‰", st.session_state.modified_content, height=300, key="edit_modified")
                    if edited_mod != st.session_state.modified_content:
                        st.session_state.modified_content = edited_mod

                adopt_col, dl_col = st.columns(2)
                with adopt_col:
                    if st.button("é‡‡ç”¨ä¿®æ”¹åç¨¿ä»¶ï¼ˆè¿›å…¥äººè¯ä¿®æ”¹ï¼‰", key="btn_adopt_audit", use_container_width=True, type="primary"):
                        st.session_state.kol_content = st.session_state.modified_content
                        st.success("å·²é‡‡ç”¨ï¼å¯è¿›å…¥ä¸‹æ–¹äººè¯ä¿®æ”¹")

                with dl_col:
                    from docx.shared import Pt
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = 'PingFang SC'
                    style.font.size = Pt(11)
                    output_name = f"é‡‡çº³åç¨¿ä»¶_{TODAY}"
                    doc.add_heading("é‡‡çº³åç¨¿ä»¶", 0)
                    for line in st.session_state.modified_content.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button("ğŸ“¥ ä¸‹è½½é‡‡çº³åç¨¿ä»¶", buf, f"{output_name}.docx",
                                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     key="dl_audit")

# ================================================================
# Part 2: äººè¯ä¿®æ”¹
# ================================================================
with st.container(border=True):
    st.markdown('<div id="part2-marker"></div>', unsafe_allow_html=True)
    st.markdown("#### Part 2 Â· äººè¯ä¿®æ”¹ï¼ˆå…­æ­¥å®¡è®¡æ³•ï¼‰")
    st.caption("AIæŒ‰ç…§å…­æ­¥å®¡è®¡æ³•å¯¹ç¨¿ä»¶è¿›è¡Œäººè¯ä¿®æ”¹ï¼šå–ç‚¹é€»è¾‘â†’ç»“æ„å®Œæ•´æ€§â†’å£å»äººè®¾â†’å…³é”®è¯ç¦è¯â†’è¯æœ¯å›å¡«â†’å†…å®¹ç»“æ„å æ¯”")

    if not st.session_state.kol_content:
        st.info("è¯·å…ˆä¸Šä¼ ç¨¿ä»¶å¹¶å®Œæˆå…«å¤§å®¡æ ¸")
    else:
        st.markdown(f'<div style="background:#fff;border-left:3px solid #4caf50;padding:8px 12px;font-size:13px;margin-bottom:10px;">å½“å‰ç¨¿ä»¶ï¼š{count_chinese(st.session_state.kol_content)} å­—</div>', unsafe_allow_html=True)

        if st.button("å¼€å§‹äººè¯ä¿®æ”¹ï¼ˆè‡ªåŠ¨å¾ªç¯è‡³å…«å¤§å®¡æ ¸å…¨é€šè¿‡ï¼‰", key="btn_renhua", use_container_width=True, type="primary"):
            max_retries = 5
            retry_count = 0
            current_content = st.session_state.kol_content
            final_result = None
            all_passed = False

            progress_bar = st.progress(0)
            status_text = st.empty()

            while retry_count < max_retries and not all_passed:
                retry_count += 1
                status_text.markdown(f"ğŸ”„ **ç¬¬ {retry_count} æ¬¡ç”Ÿæˆä¸­...**ï¼ˆæœ€å¤šå°è¯•{max_retries}æ¬¡ï¼‰")
                progress_bar.progress(retry_count / max_retries * 0.8)

                if retry_count == 1:
                    # ç¬¬ä¸€æ¬¡ï¼šä½¿ç”¨åŸå§‹prompt
                    prompt = RENHUA_PROMPT.replace("{content}", current_content)
                else:
                    # åç»­ï¼šæ ¹æ®å¤±è´¥é¡¹ç”Ÿæˆä¿®æ­£prompt
                    fix_hints = []
                    r = run_all_checks(final_result)
                    if r.get("check1", {}).get("status") != "pass":
                        fix_hints.append("- è°ƒæ•´å–ç‚¹é¡ºåºï¼šå¿…é¡»æŒ‰ é˜²æ•-æ°´è§£æŠ€æœ¯â†’è‡ªæŠ¤åŠ›â†’åŸºç¡€è¥å…» é¡ºåº")
                    if r.get("check2", {}).get("status") != "pass":
                        wc = r['check2']['count']
                        hint = "å­—æ•°ä¸è¶³ï¼Œéœ€æ‰©å……" if wc < 800 else "å­—æ•°è¶…æ ‡ï¼Œéœ€ç²¾ç®€"
                        fix_hints.append(f"- {hint}ï¼šå½“å‰{wc}å­—ï¼Œå¿…é¡»åœ¨800-900å­—ä¹‹é—´")
                    if r.get("check3", {}).get("status") != "pass":
                        fix_hints.append("- å¿…é¡»æä¾›3ä¸ªå¤‡é€‰æ ‡é¢˜")
                    if r.get("check4", {}).get("status") != "pass":
                        missing = r['check4'].get('missing', [])
                        fix_hints.append(f"- è¡¥å……æ ‡ç­¾ï¼š{', '.join(missing)}")
                    if r.get("check5", {}).get("status") != "pass":
                        fix_hints.append("- æ ‡é¢˜å¿…å«ã€é€‚åº¦æ°´è§£ã€é˜²æ•ã€ç§‘æ™®ã€‘ï¼Œæ­£æ–‡å¿…å«ã€é€‚åº¦æ°´è§£ã€é˜²æ•ã€èƒ½æ©å…¨æŠ¤ã€‘")
                    if r.get("check6", {}).get("status") != "pass":
                        found = [x['word'] for x in r['check6']['items'] if x['found']]
                        fix_hints.append(f"- åˆ é™¤ç¦è¯ï¼š{', '.join(found)}")
                    if r.get("check7", {}).get("status") != "pass":
                        fix_hints.append("- è¡¥å……æ¶¦è‰²å–ç‚¹ï¼šç¡®ä¿10ä¸ªå°æ–¹å‘æ ¸å¿ƒå†…å®¹éƒ½æœ‰ä½“ç°")
                    if r.get("check8", {}).get("status") != "pass":
                        missing8 = [x['text'] for x in r['check8']['items'] if not x['found']]
                        fix_hints.append(f"- å¿…é¡»åŸå°ä¸åŠ¨åŠ å…¥ä»¥ä¸‹è¯æœ¯ï¼š\n  " + "\n  ".join(missing8))

                    fix_text = "\n".join(fix_hints)
                    prompt = f"""è¯·ä¿®æ­£ä»¥ä¸‹ç¨¿ä»¶ï¼Œè§£å†³æ£€æµ‹åˆ°çš„é—®é¢˜ï¼š

ã€éœ€è¦ä¿®æ­£çš„é—®é¢˜ã€‘
{fix_text}

ã€å½“å‰ç¨¿ä»¶ã€‘
{final_result}

ã€ä¿®æ­£è¦æ±‚ã€‘
1. âš ï¸ æ­£æ–‡å¿…é¡»åœ¨800-900å­—ä¹‹é—´ï¼ˆæœ€é‡è¦ï¼ï¼‰
2. å¿…é¡»æä¾›3ä¸ªå¤‡é€‰æ ‡é¢˜ï¼ˆåŒ…å«ï¼šé€‚åº¦æ°´è§£ã€é˜²æ•ã€ç§‘æ™®ï¼‰
3. å¿…é¡»åŒ…å«10ä¸ªä»¥ä¸Šæ ‡ç­¾ï¼ŒåŒ…æ‹¬ï¼š#èƒ½æ©å…¨æŠ¤ #é€‚åº¦æ°´è§£ #é€‚åº¦æ°´è§£å¥¶ç²‰æ¨è #ç¬¬ä¸€å£å¥¶ç²‰
4. åˆ é™¤æ‰€æœ‰ç¦è¯ï¼ˆæ•å®ã€å¥¶ç“¶ã€å¥¶å˜´ã€æ–°ç”Ÿå„¿ã€è¿‡æ•ã€ç–¾ç—…ã€é¢„é˜²ã€å…ç–«ï¼‰
5. å¿…é¡»åŒ…å«å…¨éƒ¨10å¥ä¸å¯ä¿®æ”¹è¯æœ¯ï¼ˆå­—å­—ä¸å·®ï¼‰
6. ä¿æŒå°çº¢ä¹¦æ´»äººæ„Ÿçˆ†æ–‡é£æ ¼ï¼ˆç”¨"å§å¦¹ä»¬""çœŸçš„ç»äº†"ç­‰å£è¯­åŒ–è¡¨è¾¾ï¼‰

è¯·ç›´æ¥è¾“å‡ºä¿®æ­£åçš„å®Œæ•´ç¨¿ä»¶ï¼š
### æ ‡é¢˜å¤‡é€‰ï¼ˆ3ä¸ªï¼‰
### æ­£æ–‡ï¼ˆ800-900å­—ï¼Œå¿…é¡»å†™å¤Ÿï¼ï¼‰
### è¯é¢˜æ ‡ç­¾ï¼ˆ10ä¸ªä»¥ä¸Šï¼‰"""

                result = call_llm_api(prompt)
                if result and not result.startswith("Error"):
                    final_result = result
                    # è‡ªåŠ¨æ’å…¥ç¼ºå¤±çš„ä¸å¯ä¿®æ”¹è¯æœ¯
                    final_result, inserted_count = auto_insert_fixed_phrases(final_result)
                    if inserted_count > 0:
                        status_text.markdown(f"ğŸ“ è‡ªåŠ¨è¡¥å……äº† {inserted_count} æ¡ç¼ºå¤±è¯æœ¯")
                    # æ£€æŸ¥æ˜¯å¦å…¨éƒ¨é€šè¿‡
                    check_result = run_all_checks(final_result)
                    pass_count = sum(1 for k in ["check1","check2","check3","check4","check5","check6","check7","check8"]
                                   if check_result.get(k, {}).get("status") == "pass")
                    status_text.markdown(f"ğŸ” ç¬¬ {retry_count} æ¬¡æ£€æŸ¥ï¼šé€šè¿‡ {pass_count}/8 é¡¹")
                    if pass_count == 8:
                        all_passed = True
                else:
                    status_text.error(f"AIè°ƒç”¨å¤±è´¥: {result}")
                    break

            progress_bar.progress(1.0)
            if all_passed:
                status_text.success(f"âœ… å…«å¤§å®¡æ ¸å…¨éƒ¨é€šè¿‡ï¼ï¼ˆå…±å°è¯• {retry_count} æ¬¡ï¼‰")
                st.session_state.renhua_result = final_result
                st.session_state.recheck_results = run_all_checks(final_result)
                st.rerun()
            elif final_result:
                status_text.warning(f"âš ï¸ å·²è¾¾æœ€å¤§å°è¯•æ¬¡æ•°({max_retries}æ¬¡)ï¼Œå½“å‰ç»“æœå¯èƒ½ä»æœ‰æœªé€šè¿‡é¡¹ï¼Œå¯æ‰‹åŠ¨ç¼–è¾‘ä¿®æ­£")
                st.session_state.renhua_result = final_result
                st.session_state.recheck_results = run_all_checks(final_result)
                st.rerun()

        if st.session_state.renhua_result:
            st.markdown("---")

            # å±•ç¤ºåŸæ–‡
            with st.expander("ğŸ“„ å®¡æ ¸åç¨¿ä»¶ï¼ˆä¿®æ”¹å‰ï¼‰", expanded=False):
                orig_html = st.session_state.kol_content.replace('\n', '<br>')
                st.markdown(f'<div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:15px;font-size:13px;line-height:1.8;">{orig_html}</div>', unsafe_allow_html=True)

            # ===== å¤æ ¸æ£€æŸ¥ =====
            st.markdown("### ğŸ” å¤æ ¸æ£€æŸ¥ï¼ˆäººè¯ä¿®æ”¹åè‡ªåŠ¨éªŒè¯ï¼‰")
            result_text = st.session_state.renhua_result

            # æ£€æŸ¥æ ‡é¢˜æ•°é‡
            title_matches = re.findall(r'###\s*æ ‡é¢˜å¤‡é€‰.*?(?=###|$)', result_text, re.DOTALL)
            title_section = title_matches[0] if title_matches else ""
            title_count = len(re.findall(r'\d+\.\s*.+', title_section))

            # æ£€æŸ¥æ­£æ–‡å­—æ•°
            body_matches = re.findall(r'###\s*æ­£æ–‡.*?(?=###|$)', result_text, re.DOTALL)
            body_section = body_matches[0] if body_matches else result_text
            body_word_count = count_chinese(body_section)

            # æ£€æŸ¥æ ‡ç­¾æ•°é‡
            tags_in_result = extract_tags(result_text)

            # æ£€æŸ¥æ´»äººæ„Ÿå…³é”®è¯ï¼ˆå¢å¼ºç‰ˆï¼‰
            human_markers = ["æˆ‘", "ä½ ", "å’±", "è¯´å®è¯", "ä¸ç’ä½ è¯´", "ä¸€å¼€å§‹", "å…¶å®", "çœŸçš„", "å§å¦¹", "ç»äº†", "æ•‘å‘½", "åæ‚”"]
            human_found = sum(1 for m in human_markers if m in result_text)
            # æ£€æŸ¥emojiä½¿ç”¨
            emoji_markers = ["ğŸ’¡", "âœ¨", "ğŸ”¥", "â—", "ğŸ‘¶", "ğŸ¼", "ğŸ’ª", "â¤ï¸", "ğŸ™‹", "ğŸ˜Š"]
            emoji_found = sum(1 for e in emoji_markers if e in result_text)
            # æ£€æŸ¥æ„Ÿå¹å·ä½¿ç”¨ï¼ˆå°çº¢ä¹¦çˆ†æ–‡ç‰¹å¾ï¼‰
            exclamation_count = result_text.count("ï¼") + result_text.count("!")

            # å¤æ ¸è¡¨æ ¼
            check_items = [
                ("å®¡æ ¸2 - å­—æ•°", f"{body_word_count}å­—ï¼ˆè¦æ±‚800-900ï¼‰", "pass" if 800 <= body_word_count <= 900 else "fail"),
                ("å®¡æ ¸3 - æ ‡é¢˜æ•°é‡", f"{title_count}ä¸ªå¤‡é€‰æ ‡é¢˜", "pass" if title_count >= 3 else "fail"),
                ("å®¡æ ¸4 - æ ‡ç­¾æ•°é‡", f"{len(tags_in_result)}ä¸ªæ ‡ç­¾", "pass" if len(tags_in_result) >= 10 else "fail"),
                ("æ´»äººæ„Ÿå…³é”®è¯", f"åŒ…å«{human_found}/{len(human_markers)}ä¸ªå£è¯­åŒ–è¡¨è¾¾", "pass" if human_found >= 5 else "warn"),
                ("Emojiä½¿ç”¨", f"åŒ…å«{emoji_found}ä¸ªemoji", "pass" if emoji_found >= 3 else "warn"),
                ("çˆ†æ–‡è¯­æ°”", f"{exclamation_count}ä¸ªæ„Ÿå¹å·", "pass" if exclamation_count >= 5 else "warn"),
            ]

            rows_recheck = ""
            for name, detail, status in check_items:
                if status == "pass":
                    icon = "âœ…"
                    bg = "#f0fff4"
                elif status == "fail":
                    icon = "âŒ"
                    bg = "#fff5f5"
                else:
                    icon = "âš ï¸"
                    bg = "#fffde7"
                rows_recheck += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:8px;">{name}</td><td style="border:1px solid #ddd;padding:8px;">{detail}</td><td style="border:1px solid #ddd;padding:8px;font-weight:bold;">{icon}</td></tr>'

            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:8px 0 16px 0;">
            <thead><tr style="background:#e3f2fd;"><th style="border:1px solid #ddd;padding:8px;">æ£€æŸ¥é¡¹</th><th style="border:1px solid #ddd;padding:8px;">è¯¦æƒ…</th><th style="border:1px solid #ddd;padding:8px;">ç»“æœ</th></tr></thead>
            <tbody>{rows_recheck}</tbody></table>''', unsafe_allow_html=True)

            # æ´»äººæ„Ÿè¯¦ç»†æ£€æŸ¥
            with st.expander("ğŸ“ å°çº¢ä¹¦çˆ†æ–‡ç¬”è®°æ”»ç•¥ Â· æ´»äººæ„Ÿæ£€æŸ¥", expanded=True):
                st.markdown("**ğŸ—£ï¸ å£è¯­åŒ–è¡¨è¾¾**")
                markers_detail = []
                for m in human_markers:
                    if m in result_text:
                        markers_detail.append(f'<span style="background:#c8e6c9;padding:2px 6px;border-radius:3px;margin:2px;">âœ… {m}</span>')
                    else:
                        markers_detail.append(f'<span style="background:#ffcdd2;padding:2px 6px;border-radius:3px;margin:2px;">âŒ {m}</span>')
                st.markdown(f'<div style="line-height:2.2;">{"".join(markers_detail)}</div>', unsafe_allow_html=True)

                st.markdown("**ğŸ˜Š Emojiä½¿ç”¨**")
                emoji_detail = []
                for e in emoji_markers:
                    if e in result_text:
                        emoji_detail.append(f'<span style="background:#c8e6c9;padding:2px 6px;border-radius:3px;margin:2px;">âœ… {e}</span>')
                    else:
                        emoji_detail.append(f'<span style="background:#ffcdd2;padding:2px 6px;border-radius:3px;margin:2px;">âŒ {e}</span>')
                st.markdown(f'<div style="line-height:2.2;">{"".join(emoji_detail)}</div>', unsafe_allow_html=True)

                st.markdown(f"**ğŸ”¥ çˆ†æ–‡è¯­æ°”**ï¼šå…±{exclamation_count}ä¸ªæ„Ÿå¹å·ï¼ˆå»ºè®®â‰¥5ä¸ªï¼‰")
                st.caption("å°çº¢ä¹¦çˆ†æ–‡ç‰¹å¾ï¼šå¤šç”¨æ„Ÿå¹å·ã€emojiã€å£è¯­åŒ–è¡¨è¾¾ï¼Œåƒé—ºèœœèŠå¤©ä¸€æ ·è‡ªç„¶")

            # å±•ç¤ºAIç»“æœ
            st.markdown("---")
            st.markdown("### äººè¯ä¿®æ”¹ç»“æœ")
            st.markdown(st.session_state.renhua_result)

            # ç¼–è¾‘ + æ“ä½œ
            with st.expander("éœ€è¦å¾®è°ƒï¼Ÿç‚¹å‡»ç¼–è¾‘", expanded=False):
                edited_renhua = st.text_area("äººè¯ä¿®æ”¹å†…å®¹ï¼ˆå¯ç¼–è¾‘ï¼‰", st.session_state.renhua_result, height=400, key="edit_renhua")
                if edited_renhua != st.session_state.renhua_result:
                    st.session_state.renhua_result = edited_renhua

            if st.button("é‡‡ç”¨äººè¯ä¿®æ”¹ç»“æœ â†’ è¿›å…¥Part 3å¤æ ¸", key="btn_adopt_renhua", use_container_width=True, type="primary"):
                st.session_state.renhua_adopted = True
                st.session_state.recheck_content = st.session_state.renhua_result
                st.session_state.recheck_results = None
                st.session_state.final_ready = False
                st.success("å·²é‡‡ç”¨ï¼è¯·åœ¨ä¸‹æ–¹Part 3è¿›è¡Œå¤æ ¸æ£€æŸ¥")
                st.rerun()

# ================================================================
# Part 3: å¤æ ¸æ£€æŸ¥
# ================================================================
with st.container(border=True):
    st.markdown('<div id="part3-marker"></div>', unsafe_allow_html=True)
    st.markdown("#### Part 3 Â· å¤æ ¸æ£€æŸ¥ï¼ˆå†æ¬¡å…«å¤§å®¡æ ¸ï¼‰")
    st.caption("å¯¹äººè¯ä¿®æ”¹åçš„ç¨¿ä»¶è¿›è¡Œå…«å¤§å®¡æ ¸ï¼Œç¡®ä¿åˆè§„åå¯ç¼–è¾‘å¾®è°ƒ")

    if not st.session_state.renhua_adopted or not st.session_state.recheck_content:
        st.info("è¯·å…ˆå®ŒæˆPart 2äººè¯ä¿®æ”¹å¹¶é‡‡ç”¨ç»“æœ")
    else:
        # æ˜¾ç¤ºå½“å‰å†…å®¹å­—æ•°
        recheck_wc = count_chinese(st.session_state.recheck_content)
        st.markdown(f'<div style="background:#fff;border-left:3px solid #ff9800;padding:8px 12px;font-size:13px;margin-bottom:10px;">å¾…å¤æ ¸ç¨¿ä»¶ï¼š{recheck_wc} å­—</div>', unsafe_allow_html=True)

        # å¼€å§‹å¤æ ¸æŒ‰é’®
        if st.button("å¼€å§‹å¤æ ¸ï¼ˆå…«å¤§å®¡æ ¸ï¼‰", key="btn_recheck", use_container_width=True, type="primary"):
            st.session_state.recheck_results = run_all_checks(st.session_state.recheck_content)
            st.rerun()

        # æ˜¾ç¤ºå¤æ ¸ç»“æœ
        if st.session_state.recheck_results:
            r3 = st.session_state.recheck_results

            # ç»Ÿè®¡æ¦‚è§ˆ
            pass_count3 = sum(1 for k in ["check1","check2","check3","check4","check5","check6","check7","check8"]
                           if r3.get(k, {}).get("status") == "pass")
            fail_count3 = 8 - pass_count3

            st.markdown("### å¤æ ¸ç»“æœæ¦‚è§ˆ")
            m3_1, m3_2, m3_3 = st.columns(3)
            m3_1.metric("é€šè¿‡", f"{pass_count3}/8", delta="è‰¯å¥½" if pass_count3 >= 6 else "éœ€ä¿®æ”¹")
            m3_2.metric("éœ€ä¿®æ”¹", f"{fail_count3}")
            m3_3.metric("å­—æ•°", f"{r3['check2']['count']}å­—", delta="800-900" if 800 <= r3['check2']['count'] <= 900 else "éœ€è°ƒæ•´")

            # ç®€åŒ–çš„å®¡æ ¸ç»“æœè¡¨æ ¼
            st.markdown("### å…«å¤§å®¡æ ¸ç»“æœ")
            check_names = [
                ("check1", "å®¡æ ¸1-å–ç‚¹é¡ºåº"),
                ("check2", "å®¡æ ¸2-å­—æ•°æ£€æŸ¥"),
                ("check3", "å®¡æ ¸3-æ ‡é¢˜æ•°é‡"),
                ("check4", "å®¡æ ¸4-è¯é¢˜æ ‡ç­¾"),
                ("check5", "å®¡æ ¸5-å…³é”®è¯"),
                ("check6", "å®¡æ ¸6-ç¦è¯æ£€æŸ¥"),
                ("check7", "å®¡æ ¸7-æ¶¦è‰²å–ç‚¹"),
                ("check8", "å®¡æ ¸8-ä¸å¯ä¿®æ”¹å–ç‚¹"),
            ]
            rows3 = ""
            for key, name in check_names:
                status = r3.get(key, {}).get("status", "fail")
                icon = "âœ…" if status == "pass" else "âŒ"
                bg = "#f0fff4" if status == "pass" else "#fff5f5"
                # è¯¦æƒ…
                if key == "check2":
                    detail = f"{r3['check2']['count']}å­—ï¼ˆè¦æ±‚800-900ï¼‰"
                elif key == "check3":
                    detail = f"{r3['check3']['count']}ä¸ªæ ‡é¢˜"
                elif key == "check4":
                    detail = f"{r3['check4']['count']}ä¸ªæ ‡ç­¾ï¼Œç¼ºå¤±{len(r3['check4']['missing'])}ä¸ª"
                elif key == "check7":
                    missing7 = len([x for x in r3['check7']['items'] if not x['found']])
                    detail = f"ç¼ºå¤±{missing7}/10ä¸ªæ¶¦è‰²å–ç‚¹"
                elif key == "check8":
                    missing8 = len([x for x in r3['check8']['items'] if not x['found']])
                    detail = f"ç¼ºå¤±{missing8}/10ä¸ªä¸å¯ä¿®æ”¹å–ç‚¹"
                else:
                    detail = "â€”"
                rows3 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:8px;">{name}</td><td style="border:1px solid #ddd;padding:8px;">{detail}</td><td style="border:1px solid #ddd;padding:8px;font-weight:bold;">{icon}</td></tr>'

            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:8px 0;">
            <thead><tr style="background:#fff3e0;"><th style="border:1px solid #ddd;padding:8px;">æ£€æŸ¥é¡¹</th><th style="border:1px solid #ddd;padding:8px;">è¯¦æƒ…</th><th style="border:1px solid #ddd;padding:8px;">ç»“æœ</th></tr></thead>
            <tbody>{rows3}</tbody></table>''', unsafe_allow_html=True)

            # å¦‚æœæœ‰å¤±è´¥é¡¹ï¼Œæ˜¾ç¤ºé‡æ–°ç”ŸæˆæŒ‰é’®
            if fail_count3 > 0:
                st.markdown("---")
                st.warning(f"âš ï¸ æ£€æµ‹åˆ° {fail_count3} é¡¹æœªé€šè¿‡ï¼Œéœ€è¦é‡æ–°ç”Ÿæˆäººè¯ç‰ˆæœ¬")

                if st.button("ğŸ”„ é‡æ–°ç”Ÿæˆäººè¯ç‰ˆæœ¬ï¼ˆAIè‡ªåŠ¨ä¿®æ­£ï¼‰", key="btn_regenerate", use_container_width=True, type="primary"):
                    with st.spinner("AIé‡æ–°ç”Ÿæˆä¸­ï¼Œè‡ªåŠ¨ä¿®æ­£æœªé€šè¿‡é¡¹..."):
                        # æ„å»ºé’ˆå¯¹å¤±è´¥é¡¹çš„ä¿®æ­£æç¤º
                        fix_hints = []
                        if r3.get("check1", {}).get("status") != "pass":
                            fix_hints.append("- è°ƒæ•´å–ç‚¹é¡ºåºï¼šå¿…é¡»æŒ‰ é˜²æ•-æ°´è§£æŠ€æœ¯â†’è‡ªæŠ¤åŠ›â†’åŸºç¡€è¥å…» é¡ºåº")
                        if r3.get("check2", {}).get("status") != "pass":
                            wc3 = r3['check2']['count']
                            hint3 = "å­—æ•°ä¸è¶³ï¼Œéœ€æ‰©å……" if wc3 < 800 else "å­—æ•°è¶…æ ‡ï¼Œéœ€ç²¾ç®€"
                            fix_hints.append(f"- {hint3}ï¼šå½“å‰{wc3}å­—ï¼Œå¿…é¡»åœ¨800-900å­—ä¹‹é—´")
                        if r3.get("check3", {}).get("status") != "pass":
                            fix_hints.append("- è¡¥å……æ ‡é¢˜ï¼šå¿…é¡»æä¾›3ä¸ªå¤‡é€‰æ ‡é¢˜")
                        if r3.get("check4", {}).get("status") != "pass":
                            missing_tags = r3['check4'].get('missing', [])
                            fix_hints.append(f"- è¡¥å……æ ‡ç­¾ï¼šç¼ºå¤± {', '.join(missing_tags[:3])}...")
                        if r3.get("check6", {}).get("status") != "pass":
                            found_forbidden = [x['word'] for x in r3['check6']['items'] if x['found']]
                            fix_hints.append(f"- åˆ é™¤ç¦è¯ï¼š{', '.join(found_forbidden)}")
                        if r3.get("check7", {}).get("status") != "pass":
                            fix_hints.append("- è¡¥å……æ¶¦è‰²å–ç‚¹ï¼šç¡®ä¿10ä¸ªå°æ–¹å‘éƒ½æœ‰ä½“ç°")
                        if r3.get("check8", {}).get("status") != "pass":
                            missing_fixed = [x['text'][:20]+"..." for x in r3['check8']['items'] if not x['found']][:3]
                            fix_hints.append(f"- è¡¥å……ä¸å¯ä¿®æ”¹å–ç‚¹ï¼š{', '.join(missing_fixed)}")

                        fix_prompt = "\n".join(fix_hints)

                        regen_prompt = f"""ä½ æ˜¯å°çº¢ä¹¦çˆ†æ–‡å†™æ‰‹ã€‚è¯·ä¿®æ­£ä»¥ä¸‹ç¨¿ä»¶ï¼Œè§£å†³æ£€æµ‹åˆ°çš„é—®é¢˜ã€‚

ã€éœ€è¦ä¿®æ­£çš„é—®é¢˜ã€‘
{fix_prompt}

ã€åŸç¨¿ä»¶ã€‘
{st.session_state.recheck_content}

ã€ä¿®æ­£è¦æ±‚ - æŒ‰é‡è¦æ€§æ’åºã€‘
1. âš ï¸ æ­£æ–‡å¿…é¡»åœ¨800-900å­—ä¹‹é—´ï¼ˆæœ€é‡è¦ï¼å†™å¤Ÿå­—æ•°ï¼ï¼‰
2. å¿…é¡»æœ‰å°çº¢ä¹¦æ´»äººæ„Ÿï¼ˆç”¨"å§å¦¹ä»¬""çœŸçš„ç»äº†""è¯´å®è¯"ç­‰å£è¯­åŒ–è¡¨è¾¾ï¼‰
3. å¿…é¡»æä¾›3ä¸ªå¤‡é€‰æ ‡é¢˜ï¼ˆåŒ…å«ï¼šé€‚åº¦æ°´è§£ã€é˜²æ•ã€ç§‘æ™®ï¼‰
4. å¿…é¡»åŒ…å«10ä¸ªä»¥ä¸Šæ ‡ç­¾ï¼ˆå¿…å« #èƒ½æ©å…¨æŠ¤ #é€‚åº¦æ°´è§£ï¼‰
5. åˆ é™¤ç¦è¯ï¼ˆæ•å®ã€å¥¶ç“¶ã€å¥¶å˜´ã€æ–°ç”Ÿå„¿ã€è¿‡æ•ã€ç–¾ç—…ã€é¢„é˜²ã€å…ç–«ï¼‰
6. å¿…é¡»åŒ…å«å…¨éƒ¨10å¥ä¸å¯ä¿®æ”¹è¯æœ¯ï¼ˆå­—å­—ä¸å·®ï¼‰

è¯·ç›´æ¥è¾“å‡ºä¿®æ­£åçš„å®Œæ•´ç¨¿ä»¶ï¼š
### æ ‡é¢˜å¤‡é€‰ï¼ˆ3ä¸ªï¼‰
### æ­£æ–‡ï¼ˆ800-900å­—ï¼Œå¿…é¡»å†™å¤Ÿï¼ï¼‰
### è¯é¢˜æ ‡ç­¾ï¼ˆ10ä¸ªä»¥ä¸Šï¼‰"""

                        result = call_llm_api(regen_prompt)
                        if result and not result.startswith("Error"):
                            # è‡ªåŠ¨æ’å…¥ç¼ºå¤±çš„ä¸å¯ä¿®æ”¹è¯æœ¯
                            result, inserted_count = auto_insert_fixed_phrases(result)
                            if inserted_count > 0:
                                st.info(f"ğŸ“ è‡ªåŠ¨è¡¥å……äº† {inserted_count} æ¡ç¼ºå¤±è¯æœ¯")
                            st.session_state.recheck_content = result
                            # åŒæ­¥text_areaçš„key
                            st.session_state.edit_recheck_content = result
                            st.session_state.recheck_results = run_all_checks(result)
                            st.rerun()
                        else:
                            st.error(f"AIè°ƒç”¨å¤±è´¥: {result}")

                # ä¹Ÿå…è®¸æ‰‹åŠ¨ç¼–è¾‘
                st.markdown("---")
                st.markdown("### æˆ–æ‰‹åŠ¨ç¼–è¾‘ä¿®æ­£")
            else:
                st.success("ğŸ‰ æ­å–œï¼å…«å¤§å®¡æ ¸å…¨éƒ¨é€šè¿‡ï¼")
                st.markdown("---")
                st.markdown("### æœ€ç»ˆç¨¿ä»¶é¢„è§ˆ")

            # å¯ç¼–è¾‘çš„æ­£æ–‡åŒºåŸŸ
            edited_recheck = st.text_area(
                "ç¼–è¾‘æ­£æ–‡å†…å®¹",
                st.session_state.recheck_content,
                height=400,
                key="edit_recheck_content"
            )
            if edited_recheck != st.session_state.recheck_content:
                st.session_state.recheck_content = edited_recheck

            # å®æ—¶å­—æ•°æ˜¾ç¤º
            current_wc = count_chinese(edited_recheck)
            wc_color = "#4caf50" if 800 <= current_wc <= 900 else "#f44336"
            st.markdown(f'<div style="text-align:right;color:{wc_color};font-weight:bold;">å½“å‰å­—æ•°ï¼š{current_wc}/900</div>', unsafe_allow_html=True)

            # æ‰‹åŠ¨ä¿®æ”¹åé‡æ–°æ£€æŸ¥æŒ‰é’®
            if fail_count3 > 0:
                if st.button("ğŸ” é‡æ–°æ£€æŸ¥ï¼ˆæ‰‹åŠ¨ä¿®æ”¹åï¼‰", key="btn_manual_recheck", use_container_width=True):
                    st.session_state.recheck_results = run_all_checks(st.session_state.recheck_content)
                    st.rerun()

            # åªæœ‰å…¨éƒ¨é€šè¿‡æ‰èƒ½è¿›å…¥Part 4
            if fail_count3 == 0:
                if st.button("âœ… ç¡®è®¤å¤æ ¸å®Œæˆ â†’ è¿›å…¥Part 4ç»ˆç¨¿", key="btn_confirm_recheck", use_container_width=True, type="primary"):
                    st.session_state.final_content = st.session_state.recheck_content
                    st.session_state.final_ready = True
                    st.success("å¤æ ¸å®Œæˆï¼è¯·åœ¨Part 4é¢„è§ˆå¹¶ä¸‹è½½ç»ˆç¨¿")
                    st.rerun()
            else:
                st.info("ğŸ’¡ è¯·å…ˆä¿®æ­£æ‰€æœ‰æœªé€šè¿‡é¡¹ï¼Œå…«å¤§å®¡æ ¸å…¨éƒ¨é€šè¿‡åæ‰èƒ½è¿›å…¥Part 4")

# ================================================================
# Part 4: ç»ˆç¨¿å®Œæˆ
# ================================================================
with st.container(border=True):
    st.markdown('<div id="part4-marker"></div>', unsafe_allow_html=True)
    st.markdown("#### Part 4 Â· ç»ˆç¨¿å®Œæˆ")
    st.caption("é¢„è§ˆç»ˆç¨¿å¹¶ä¸‹è½½")

    if not st.session_state.final_ready or not st.session_state.final_content:
        st.info("è¯·å…ˆå®ŒæˆPart 3å¤æ ¸æ£€æŸ¥")
    else:
        # ç»ˆç¨¿ä¿¡æ¯
        final_wc = count_chinese(st.session_state.final_content)
        final_tags = extract_tags(st.session_state.final_content)
        dir_name = st.session_state.selected_direction if st.session_state.selected_direction != DIRECTION_OPTIONS[0] else "æœªæŒ‡å®šæ–¹å‘"

        # ä¿¡æ¯å¡ç‰‡
        st.markdown(f'''
        <div style="background:#fff;border:2px solid #f48fb1;border-radius:10px;padding:15px;margin-bottom:15px;">
            <div style="font-size:16px;font-weight:bold;color:#c2185b;margin-bottom:10px;">ğŸ“‹ ç»ˆç¨¿ä¿¡æ¯</div>
            <div style="display:flex;gap:20px;flex-wrap:wrap;">
                <div>ğŸ“ å­—æ•°ï¼š<b>{final_wc}</b></div>
                <div>ğŸ·ï¸ æ ‡ç­¾ï¼š<b>{len(final_tags)}ä¸ª</b></div>
                <div>ğŸ“‚ æ–¹å‘ï¼š<b>{dir_name}</b></div>
                <div>ğŸ“… æ—¥æœŸï¼š<b>{TODAY}</b></div>
            </div>
        </div>
        ''', unsafe_allow_html=True)

        # ç»ˆç¨¿é¢„è§ˆ
        st.markdown("### ç»ˆç¨¿é¢„è§ˆ")
        final_html = st.session_state.final_content.replace('\n', '<br>')
        st.markdown(f'''<div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:20px;font-size:14px;line-height:2.0;max-height:500px;overflow-y:auto;">
        {final_html}
        </div>''', unsafe_allow_html=True)

        # ä¸‹è½½æŒ‰é’®
        st.markdown("---")
        dl_col1, dl_col2 = st.columns(2)

        with dl_col1:
            from docx.shared import Pt, RGBColor
            doc_final = Document()
            style_final = doc_final.styles['Normal']
            style_final.font.name = 'PingFang SC'
            style_final.font.size = Pt(11)
            output_name_final = f"KOL_{TODAY}_ç»ˆç¨¿"
            doc_final.add_heading(output_name_final, 0)
            doc_final.add_paragraph(f"ç”Ÿæˆæ—¶é—´: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc_final.add_paragraph(f"ç¨¿ä»¶æ–¹å‘: {dir_name}")
            doc_final.add_paragraph(f"å­—æ•°: {final_wc}")
            doc_final.add_paragraph("â”€" * 50)
            for line in st.session_state.final_content.split('\n'):
                if line.strip():
                    doc_final.add_paragraph(line.strip())
            buf_final = io.BytesIO()
            doc_final.save(buf_final)
            buf_final.seek(0)
            st.download_button(
                "ğŸ“¥ ä¸‹è½½ç»ˆç¨¿ (.docx)",
                buf_final,
                f"{output_name_final}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_final",
                use_container_width=True
            )

        with dl_col2:
            # çº¯æ–‡æœ¬ä¸‹è½½
            st.download_button(
                "ğŸ“„ ä¸‹è½½çº¯æ–‡æœ¬ (.txt)",
                st.session_state.final_content,
                f"KOL_{TODAY}_ç»ˆç¨¿.txt",
                "text/plain",
                key="dl_final_txt",
                use_container_width=True
            )

        st.success("ğŸ‰ æ­å–œï¼ç»ˆç¨¿å·²å®Œæˆï¼Œå¯ä¸‹è½½ä½¿ç”¨")

# ========== Footer ==========
st.markdown("---")
dir_label = st.session_state.selected_direction if st.session_state.selected_direction != DIRECTION_OPTIONS[0] else "èƒ½æ©å…¨æŠ¤"
st.caption(f"ğŸ¤– èµæ„AIå®¡ç¨¿ç³»ç»Ÿ v4.0 Â· {dir_label}")
